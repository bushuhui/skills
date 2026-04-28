#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Markdown to DOCX Converter with Template Style Matching

使用 python-docx 将 Markdown 转换为 Word DOCX，匹配参考模板的样式。
适用于需要严格格式控制的文档（公文、技术方案、报告等）。

用法：
    python md2docx_template.py <input.md> --template <template.docx> -o <output.docx>
"""

import sys
import os
import re
import argparse
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from typing import List, Tuple


# ============================================================
# 样式配置
# ============================================================

STYLE_CONFIG = {
    'Title': {
        'font_name': '黑体',
        'font_size': Pt(16),
        'bold': True,
        'style': 'Normal',  # 文档标题用 Normal 样式 + 黑体加粗
    },
    'H1': {
        'font_name': '黑体',
        'font_size': Pt(16),
        'bold': True,
        'style': 'Heading 1',
    },
    'H2': {
        'font_name': '黑体',
        'font_size': Pt(14),
        'bold': True,
        'style': 'Heading 2',
    },
    'Normal': {
        'font_name': '宋体',
        'font_size': Pt(12),
        'bold': False,
        'style': 'Normal',
    },
}

# Heading 1 关键词（大章节标题）
H1_KEYWORDS = ['研究内容', '主要性能指标', '研究目标', '技术路线', '研究方法',
               '研究方案', '预期成果', '研究背景', '项目概述', '总体方案']


# ============================================================
# Markdown 解析器
# ============================================================

class MarkdownParser:
    """Markdown 解析器，将 Markdown 文本解析为结构化条目"""

    @staticmethod
    def parse(md_content: str) -> List[Tuple[str, str]]:
        """
        解析 Markdown 内容为 (style_id, text) 列表

        Args:
            md_content: Markdown 文件内容

        Returns:
            entries 列表：[(style_id, text), ...]
            style_id 可选值: 'Title', 'H1', 'H2', 'Normal'
        """
        lines = md_content.strip().split('\n')
        entries = []

        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue

            # H2 标题
            if line.startswith('## '):
                title = line[3:].strip()
                if any(kw in title for kw in H1_KEYWORDS):
                    entries.append(('H1', title))
                else:
                    entries.append(('H2', title))
                i += 1

            # 文档标题（第一个 # 行）
            elif line.startswith('# '):
                entries.append(('Title', line[2:].strip()))
                i += 1

            # 编号行（如技术指标列表）— 独立成段
            elif re.match(r'^\d+\.', line):
                entries.append(('Normal', line))
                i += 1

            # 普通文本段落 — 合并连续行
            else:
                parts = []
                while i < len(lines):
                    l = lines[i].strip()
                    if not l or l.startswith('#') or re.match(r'^\d+\.', l):
                        break
                    parts.append(l)
                    i += 1
                if parts:
                    # 合并时不加额外空格，因为中文不需要
                    entries.append(('Normal', ''.join(parts)))

        return entries


# ============================================================
# DOCX 构建器
# ============================================================

class DocxBuilder:
    """使用模板样式构建 DOCX 文档"""

    def __init__(self, template_path: str):
        self.doc = Document(template_path)
        # 清空模板原有内容
        for p in self.doc.paragraphs:
            p.clear()

    def build(self, entries: List[Tuple[str, str]]) -> None:
        """
        根据解析的条目构建文档

        Args:
            entries: (style_id, text) 列表
        """
        for idx, (style_id, text) in enumerate(entries):
            if idx < len(self.doc.paragraphs):
                p = self.doc.paragraphs[idx]
            else:
                p = self.doc.add_paragraph()

            p.clear()
            run = p.add_run(text)

            config = STYLE_CONFIG[style_id]
            p.style = self.doc.styles[config['style']]
            run.font.name = config['font_name']
            run.font.size = config['font_size']
            run.font.bold = config['bold']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), config['font_name'])

        # 删除多余的段落
        while len(self.doc.paragraphs) > len(entries):
            p = self.doc.paragraphs[-1]
            p._element.getparent().remove(p._element)

    def save(self, output_path: str) -> str:
        """保存文档"""
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        self.doc.save(output_path)
        return output_path


# ============================================================
# 主流程
# ============================================================

def convert(md_path: str, template_path: str, output_path: str) -> str:
    """
    完整转换流程

    Args:
        md_path: Markdown 文件路径
        template_path: 参考模板 DOCX 路径
        output_path: 输出 DOCX 路径

    Returns:
        输出文件路径
    """
    # 1. 读取 Markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # 2. 解析
    parser = MarkdownParser()
    entries = parser.parse(md_content)
    print(f"解析到 {len(entries)} 个段落")

    # 3. 构建文档
    builder = DocxBuilder(template_path)
    builder.build(entries)

    # 4. 保存
    output_path = builder.save(output_path)
    return output_path


def main():
    """命令行入口"""
    parser = argparse.ArgumentParser(description='Markdown to DOCX 转换器（模板样式匹配）')
    parser.add_argument('input', help='输入 Markdown 文件路径')
    parser.add_argument('--template', '-t', required=True, help='参考模板 DOCX 文件路径')
    parser.add_argument('--output', '-o', default=None, help='输出 DOCX 文件路径（默认与输入同名）')

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"错误：文件不存在：{args.input}")
        sys.exit(1)

    if not os.path.exists(args.template):
        print(f"错误：模板文件不存在：{args.template}")
        sys.exit(1)

    output_path = args.output
    if not output_path:
        base, _ = os.path.splitext(args.input)
        output_path = base + '.docx'

    result = convert(args.input, args.template, output_path)
    print(f"✅ 转换完成：{result}")
    print(f"   参考模板：{args.template}")


if __name__ == '__main__':
    main()
