import json
import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"
INPUT_DOCX = TESTS_DIR / "测试文档.docx"

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "ep": "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties",
}

FRONT_HEADINGS = {"摘要", "Abstract", "目录", "图目录", "表目录", "参考文献", "致谢", "附录"}
EXCLUDED_STYLES = {"table of figures", "图目录2", "表目录3", "表目录2", "图目录"}
STYLE_HINTS = {"Subtitle", "Title", "1正文", "heading 1", "heading 2", "heading 3"}

NUMERIC_HEADING_PATTERNS = [
    (re.compile(r"^(\d+\.\d+\.\d+)\s*(.+)$"), 3),
    (re.compile(r"^(\d+\.\d+)\s*(.+)$"), 2),
    (re.compile(r"^(\d+)\s+(.+)$"), 1),
    (re.compile(r"^(第[一二三四五六七八九十百千0-9]+章)\s*(.+)$"), 1),
    (re.compile(r"^(第[一二三四五六七八九十百千0-9]+节)\s*(.+)$"), 2),
    (re.compile(r"^([一二三四五六七八九十]+、)\s*(.+)$"), 2),
    (re.compile(r"^(（[一二三四五六七八九十]+）)\s*(.+)$"), 3),
]
CAPTION_PATTERNS = [
    re.compile(r"^[图表]\s*\d+(\.\d+)*"),
    re.compile(r"^(Fig\.?|Figure)\s*\d+", re.IGNORECASE),
    re.compile(r"^Table\s*\d+", re.IGNORECASE),
]


def first_or_none(values):
    return values[0] if values else None


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def load_extended_properties(zip_file: ZipFile) -> dict[str, object]:
    if "docProps/app.xml" not in zip_file.namelist():
        return {}

    root = etree.fromstring(zip_file.read("docProps/app.xml"))
    result: dict[str, object] = {}
    for child in root:
        key = local_name(child.tag)
        text = child.text.strip() if child.text else ""
        if text.isdigit():
            result[key] = int(text)
        else:
            result[key] = text
    return result


def load_style_metadata(zip_file: ZipFile) -> dict[str, dict[str, int | str | None]]:
    style_map: dict[str, dict[str, int | str | None]] = {}
    if "word/styles.xml" not in zip_file.namelist():
        return style_map

    styles_root = etree.fromstring(zip_file.read("word/styles.xml"))
    for style in styles_root.xpath(".//w:style", namespaces=NS):
        style_id = style.get(f"{{{NS['w']}}}styleId")
        if not style_id:
            continue
        style_name = first_or_none(style.xpath("./w:name/@w:val", namespaces=NS))
        outline_level = first_or_none(style.xpath("./w:pPr/w:outlineLvl/@w:val", namespaces=NS))
        style_map[style_id] = {
            "name": style_name,
            "outline_level": int(outline_level) if outline_level is not None else None,
        }
    return style_map


def extract_paragraph_metadata(docx_path: Path) -> tuple[Document, list[dict[str, object]]]:
    document = Document(docx_path)
    with ZipFile(docx_path) as zip_file:
        style_map = load_style_metadata(zip_file)
        document_root = etree.fromstring(zip_file.read("word/document.xml"))

    body = document_root.find("w:body", namespaces=NS)
    paragraph_elements = [child for child in body if local_name(child.tag) == "p"]

    metadata: list[dict[str, object]] = []
    for index, (paragraph, element) in enumerate(zip(document.paragraphs, paragraph_elements), start=1):
        style_id = first_or_none(element.xpath("./w:pPr/w:pStyle/@w:val", namespaces=NS))
        direct_outline = first_or_none(element.xpath("./w:pPr/w:outlineLvl/@w:val", namespaces=NS))
        style_outline = style_map.get(style_id, {}).get("outline_level")

        metadata.append(
            {
                "paragraph_index": index,
                "text": normalize_text(paragraph.text),
                "style_id": style_id,
                "style_name": paragraph.style.name if paragraph.style is not None else style_map.get(style_id, {}).get("name", ""),
                "outline_level": int(direct_outline) if direct_outline is not None else style_outline,
                "paragraph": paragraph,
            }
        )
    return document, metadata


def is_bold_paragraph(paragraph) -> bool:
    for run in paragraph.runs:
        if run.text.strip() and run.bold is True:
            return True
    return bool(paragraph.style and paragraph.style.font and paragraph.style.font.bold)


def is_caption_or_catalog(text: str, style_name: str) -> bool:
    if style_name in EXCLUDED_STYLES:
        return True
    return any(pattern.match(text) for pattern in CAPTION_PATTERNS)


def extract_heading_number(text: str) -> tuple[str | None, int | None]:
    for pattern, level in NUMERIC_HEADING_PATTERNS:
        match = pattern.match(text)
        if match:
            return match.group(1), level
    return None, None


def detect_heading(meta: dict[str, object]) -> dict[str, object]:
    text = str(meta["text"])
    style_name = str(meta["style_name"] or "")
    paragraph = meta["paragraph"]
    outline_level = meta.get("outline_level")

    if not text or is_caption_or_catalog(text, style_name):
        return {"is_heading": False, "heading_level": None}

    if text in FRONT_HEADINGS:
        return {"is_heading": True, "heading_level": 1}

    heading_number, number_level = extract_heading_number(text)
    style_hint = style_name in STYLE_HINTS
    short_text = len(text) <= 80 and not text.endswith("。")

    if heading_number and number_level is not None and (style_hint or outline_level is not None or short_text):
        return {"is_heading": True, "heading_level": min(number_level, 3)}

    if outline_level is not None and style_hint and short_text:
        return {"is_heading": True, "heading_level": min(int(outline_level) + 1, 3)}

    if style_name == "Title" and short_text:
        return {"is_heading": True, "heading_level": 3}

    if style_name in {"Subtitle", "1正文"} and short_text:
        return {"is_heading": True, "heading_level": 1}

    if is_bold_paragraph(paragraph) and short_text:
        return {"is_heading": True, "heading_level": 1}

    return {"is_heading": False, "heading_level": None}


def collect_table_text(document: Document) -> str:
    parts: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = normalize_text(cell.text)
                if text:
                    parts.append(text)
    return "\n".join(parts)


def count_comments(zip_file: ZipFile) -> int:
    if "word/comments.xml" not in zip_file.namelist():
        return 0
    root = etree.fromstring(zip_file.read("word/comments.xml"))
    return len(root.xpath(".//w:comment", namespaces=NS))


def build_program_stats(document: Document, metadata: list[dict[str, object]], zip_file: ZipFile) -> dict[str, object]:
    paragraph_texts = [item["text"] for item in metadata if item["text"]]
    full_text = "\n".join(paragraph_texts)
    table_text = collect_table_text(document)
    combined_text = full_text + ("\n" + table_text if table_text else "")

    headings = [detect_heading(item) for item in metadata]
    heading_levels = [item["heading_level"] for item in headings if item["is_heading"]]

    media_files = [name for name in zip_file.namelist() if name.startswith("word/media/")]
    comments_count = count_comments(zip_file)

    chinese_chars = re.findall(r"[\u4e00-\u9fff]", combined_text)
    english_words = re.findall(r"[A-Za-z]+(?:[-'][A-Za-z]+)*", combined_text)
    digits = re.findall(r"\d", combined_text)

    return {
        "main_body_paragraph_count": len(document.paragraphs),
        "non_empty_paragraph_count": len(paragraph_texts),
        "heading_count": len(heading_levels),
        "heading_level_1_count": heading_levels.count(1),
        "heading_level_2_count": heading_levels.count(2),
        "heading_level_3_count": heading_levels.count(3),
        "table_count": len(document.tables),
        "image_file_count": len(media_files),
        "comment_count": comments_count,
        "main_body_text_char_count_with_spaces": len(full_text),
        "main_body_text_char_count_without_spaces": len(re.sub(r"\s+", "", full_text)),
        "combined_text_char_count_with_spaces": len(combined_text),
        "combined_text_char_count_without_spaces": len(re.sub(r"\s+", "", combined_text)),
        "table_text_char_count_without_spaces": len(re.sub(r"\s+", "", table_text)),
        "chinese_character_count": len(chinese_chars),
        "english_word_count": len(english_words),
        "digit_count": len(digits),
    }


def build_report_text(result: dict[str, object]) -> str:
    lines = [
        "文档统计信息报告",
        f"输入文件: {result['input_file']}",
        "",
        "Word 内建统计:",
    ]
    for key, value in result["word_builtin_statistics"].items():
        lines.append(f"  {key}: {value}")

    lines.append("")
    lines.append("程序重算统计:")
    for key, value in result["program_statistics"].items():
        lines.append(f"  {key}: {value}")

    return "\n".join(lines)


def main() -> None:
    if not INPUT_DOCX.exists():
        raise FileNotFoundError(f"未找到输入文件: {INPUT_DOCX}")

    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    document, metadata = extract_paragraph_metadata(INPUT_DOCX)

    with ZipFile(INPUT_DOCX) as zip_file:
        result = {
            "input_file": str(INPUT_DOCX),
            "word_builtin_statistics": load_extended_properties(zip_file),
            "program_statistics": build_program_stats(document, metadata, zip_file),
        }

    json_path = TESTS_DIR / "3.2_文档统计信息.json"
    txt_path = TESTS_DIR / "3.2_文档统计信息.txt"

    json_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    txt_path.write_text(build_report_text(result), encoding="utf-8")

    print(f"[3.2] 输入文档: {INPUT_DOCX}")
    print(f"[3.2] 已输出文档统计 JSON: {json_path}")
    print(f"[3.2] 已输出文档统计 TXT: {txt_path}")
    print(f"[3.2] 程序统计段落数: {result['program_statistics']['main_body_paragraph_count']}，图片数: {result['program_statistics']['image_file_count']}")


if __name__ == "__main__":
    main()
