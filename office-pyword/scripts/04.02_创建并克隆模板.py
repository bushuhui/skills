from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from _06_docx_styles import prepare_common_styles


TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"
TEMPLATE_SOURCE = TESTS_DIR / "4.2_论文模板源.docx"
CLONE_OUTPUT = TESTS_DIR / "4.2_模板克隆结果.docx"


def build_template(source_path: Path) -> Path:
    document = Document()
    prepare_common_styles(document)
    document.settings.odd_and_even_pages_header_footer = True

    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    header = section.header.paragraphs[0]
    header.text = "学校名称：{{SCHOOL_NAME}}"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer.paragraphs[0]
    footer.text = "论文模板源文档"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("{{TITLE}}", style="ThesisTitle1")
    document.add_paragraph("作者：{{AUTHOR}}", style="ThesisBody")
    document.add_paragraph("学院：{{SCHOOL_NAME}}", style="ThesisBody")
    document.add_paragraph("关键词：{{KEYWORDS}}", style="ThesisBody")
    document.add_paragraph("摘要", style="ThesisTitle2")
    document.add_paragraph("{{ABSTRACT}}", style="ThesisBody")

    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "内容"
    table.cell(1, 0).text = "项目名称"
    table.cell(1, 1).text = "{{PROJECT_NAME}}"
    table.cell(2, 0).text = "指导教师"
    table.cell(2, 1).text = "{{SUPERVISOR}}"

    source_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(source_path)
    return source_path


def clone_template(source: Path, target: Path) -> Path:
    shutil.copyfile(source, target)
    return target


def main() -> None:
    template_path = build_template(TEMPLATE_SOURCE)
    clone_path = clone_template(template_path, CLONE_OUTPUT)
    print(f"[4.2] 已生成模板源文档: {template_path}")
    print(f"[4.2] 已完成模板克隆: {clone_path}")


if __name__ == "__main__":
    main()

