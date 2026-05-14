from __future__ import annotations

from docx import Document
from docx.shared import Pt

import _09_format_presets as PRESETS
from _12_docx_format_helpers import ensure_tests_dir, set_default_style, set_run_font

FONT_DEMOS = [
    ("宋体", "小二", "宋体中文 + Times New Roman English 123"),
    ("仿宋", "三号", "仿宋中文 + Times New Roman English 123"),
    ("Times New Roman", "四号", "Times New Roman 中英统一示例 123"),
    ("宋体", "小四", "正文常用小四字号示例，适合中文正文与英文缩写 ABCD 混排。"),
    ("仿宋", "五号", "五号字体示例，可用于图表说明、附注或其他较小字号内容。"),
]


def add_font_demo(document: Document, font_name: str, font_size, text: str) -> None:
    paragraph = document.add_paragraph()
    label_run = paragraph.add_run(f"字体={font_name}，字号={font_size}：")
    set_run_font(label_run, "黑体", "四号")
    text_run = paragraph.add_run(text)
    set_run_font(text_run, font_name, font_size)
    paragraph.paragraph_format.space_after = Pt(6)


def add_size_matrix(document: Document) -> None:
    paragraph = document.add_paragraph()
    title_run = paragraph.add_run("字号矩阵：")
    set_run_font(title_run, "黑体", "四号")
    for size_name in ("小二", "三号", "四号", "小四", "五号"):
        run = paragraph.add_run(f" [{size_name}] ")
        set_run_font(run, "宋体", size_name)


def main() -> None:
    output_dir = ensure_tests_dir(__file__)
    output_path = output_dir / "5.2_字体字号示例.docx"

    document = Document()
    set_default_style(document)

    heading = document.add_paragraph()
    heading_run = heading.add_run("5.2 字体、字号与字族示例")
    set_run_font(heading_run, "黑体", "小二")
    heading.alignment = PRESETS.resolve_alignment("居中")
    heading.paragraph_format.space_after = Pt(10)

    intro = document.add_paragraph()
    intro_run = intro.add_run("本示例重点演示宋体、仿宋、Times New Roman 以及常见字号预设在中英混排中的写入方式。")
    set_run_font(intro_run, "宋体", "小四")

    for font_name, font_size, text in FONT_DEMOS:
        add_font_demo(document, font_name, font_size, text)

    add_size_matrix(document)

    document.save(output_path)
    print(f"[5.2] 已生成字体字号示例文档: {output_path}")


if __name__ == "__main__":
    main()
