from __future__ import annotations

from docx import Document
from docx.shared import Pt

import _09_format_presets as PRESETS
from _12_docx_format_helpers import (
    apply_direct_format,
    ensure_tests_dir,
    set_default_style,
    set_paragraph_border,
    set_run_border,
    set_run_font,
)


def main() -> None:
    output_dir = ensure_tests_dir(__file__)
    output_path = output_dir / "5.3_直接格式和边框示例.docx"

    document = Document()
    set_default_style(document)

    heading = document.add_paragraph()
    heading_run = heading.add_run("5.3 直接格式和边框示例")
    set_run_font(heading_run, "黑体", "小二")
    heading.alignment = PRESETS.resolve_alignment("居中")
    heading.paragraph_format.space_after = Pt(12)

    paragraph_1 = document.add_paragraph()
    run = paragraph_1.add_run("直接格式示例：")
    apply_direct_format(run, font_name="黑体", font_size="四号", bold=True)
    run = paragraph_1.add_run("粗体")
    apply_direct_format(run, font_name="宋体", font_size="小四", bold=True)
    paragraph_1.add_run("、")
    run = paragraph_1.add_run("斜体")
    apply_direct_format(run, font_name="宋体", font_size="小四", italic=True)
    paragraph_1.add_run("、")
    run = paragraph_1.add_run("下划线")
    apply_direct_format(run, font_name="宋体", font_size="小四", underline=True)
    paragraph_1.add_run("、")
    run = paragraph_1.add_run("红色字体")
    apply_direct_format(run, font_name="宋体", font_size="小四", color_hex="C00000")
    paragraph_1.paragraph_format.space_after = Pt(10)

    paragraph_2 = document.add_paragraph()
    label_run = paragraph_2.add_run("文本边框示例：")
    apply_direct_format(label_run, font_name="黑体", font_size="四号", bold=True)
    boxed_run = paragraph_2.add_run("该片段带有字符边框")
    apply_direct_format(boxed_run, font_name="仿宋", font_size="小四", color_hex="1F497D")
    set_run_border(boxed_run, value="single", color="1F497D", size=12, space=2)
    paragraph_2.add_run("，其余文字保持普通正文格式。")
    paragraph_2.paragraph_format.space_after = Pt(10)

    paragraph_3 = document.add_paragraph()
    note_run = paragraph_3.add_run("段落边框示例：这是一段带有整段边框的说明文字，可用于提示、摘要块或关键说明。")
    apply_direct_format(note_run, font_name="宋体", font_size="小四")
    paragraph_3.paragraph_format.space_before = Pt(6)
    paragraph_3.paragraph_format.space_after = Pt(6)
    paragraph_3.paragraph_format.left_indent = Pt(12)
    paragraph_3.paragraph_format.right_indent = Pt(12)
    set_paragraph_border(paragraph_3, value="single", color="4F81BD", size=12, space=8)

    document.save(output_path)
    print(f"[5.3] 已生成直接格式和边框示例文档: {output_path}")


if __name__ == "__main__":
    main()
