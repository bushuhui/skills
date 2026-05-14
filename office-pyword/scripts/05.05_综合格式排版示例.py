from __future__ import annotations

from docx import Document

import _09_format_presets as PRESETS
from _12_docx_format_helpers import (
    apply_direct_format,
    apply_first_line_chars,
    apply_lines_spacing_before_after,
    apply_multiple_spacing,
    ensure_tests_dir,
    format_text_fragment,
    set_default_style,
    set_paragraph_border,
    set_run_border,
    set_run_font,
)


def add_title(document: Document) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run("三维重建技术格式综合示例")
    set_run_font(run, "黑体", "小二")
    paragraph.alignment = PRESETS.resolve_alignment("居中")
    paragraph.paragraph_format.space_after = PRESETS.resolve_font_size_pt("五号")
    apply_multiple_spacing(paragraph, 1.25)


def add_abstract_box(document: Document) -> None:
    paragraph = document.add_paragraph()
    label_run = paragraph.add_run("摘要：")
    apply_direct_format(label_run, font_name="黑体", font_size="四号", bold=True)
    body_run = paragraph.add_run("本示例集中演示段落格式、字体字号、字符边框、段落边框和文本片段格式化能力，便于后续继续扩展论文排版脚本。")
    apply_direct_format(body_run, font_name="仿宋", font_size="小四")
    apply_multiple_spacing(paragraph, 1.25)
    apply_lines_spacing_before_after(paragraph, before_lines=0.5, after_lines=0.5)
    set_paragraph_border(paragraph, value="single", color="4F81BD", size=12, space=6)


def add_body_paragraph(document: Document, text: str):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    apply_direct_format(run, font_name="宋体", font_size="小四")
    paragraph.alignment = PRESETS.resolve_alignment("两端对齐")
    apply_multiple_spacing(paragraph, 1.5)
    apply_lines_spacing_before_after(paragraph, before_lines=0, after_lines=0.5)
    apply_first_line_chars(paragraph, 2)
    return paragraph


def add_callout(document: Document) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run("关键结论")
    apply_direct_format(run, font_name="黑体", font_size="四号", bold=True, color_hex="C00000")
    set_run_border(run, value="single", color="C00000", size=12, space=2)
    suffix = paragraph.add_run("：只要同时写入 run.font 和 w:rFonts，就能更稳定地处理中英文混排。")
    apply_direct_format(suffix, font_name="宋体", font_size="小四")
    paragraph.alignment = PRESETS.resolve_alignment("左对齐")
    apply_multiple_spacing(paragraph, 1.25)


def main() -> None:
    output_dir = ensure_tests_dir(__file__)
    output_path = output_dir / "5.5_综合格式示例.docx"

    document = Document()
    set_default_style(document)

    add_title(document)
    add_abstract_box(document)

    paragraph = add_body_paragraph(
        document,
        "三维重建与数字孪生正在成为空间计算的重要基础能力，规范的排版脚本可以显著提升报告输出的一致性和表达质量。",
    )
    format_text_fragment(
        paragraph,
        "三维重建",
        formatter=lambda run: apply_direct_format(run, font_name="仿宋", font_size="小四", bold=True, color_hex="C00000"),
    )
    format_text_fragment(
        paragraph,
        "数字孪生",
        formatter=lambda run: apply_direct_format(run, font_name="仿宋", font_size="小四", bold=True, underline=True, color_hex="7030A0"),
    )

    paragraph = add_body_paragraph(
        document,
        "通过把段落设置、字体字号、直接格式和片段级格式化能力拆成清晰脚本，可以更容易地对现有文档进行自动化修改和验证。",
    )
    format_text_fragment(
        paragraph,
        "片段级格式化",
        formatter=lambda run: (
            apply_direct_format(run, font_name="宋体", font_size="小四", bold=True, color_hex="1F497D"),
            set_run_border(run, value="single", color="1F497D", size=12, space=2),
        ),
    )

    add_callout(document)

    document.save(output_path)
    print(f"[5.5] 已生成综合格式示例文档: {output_path}")


if __name__ == "__main__":
    main()
