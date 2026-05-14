from __future__ import annotations

from pathlib import Path

from docx import Document

from _06_docx_styles import SECTION6, apply_style_rule, get_or_add_paragraph_style, set_default_style

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def add_demo_content(document: Document) -> None:
    document.add_paragraph("样式系统示例", style=SECTION6.STYLE_NAMES["custom_lead"])
    document.add_paragraph(
        "这段内容使用自定义引用样式，演示在段落样式中集中管理字体、斜体、缩进和段前段后间距。",
        style=SECTION6.STYLE_NAMES["custom_quote"],
    )
    document.add_paragraph(
        "说明：自定义样式一旦写入文档，后续插入段落时只需要指定 style 名称，而不必重复设置同一批格式属性。",
        style=SECTION6.STYLE_NAMES["custom_note"],
    )


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = TESTS_DIR / "6.1_自定义样式示例.docx"

    document = Document()
    set_default_style(document)

    body_style = get_or_add_paragraph_style(document, SECTION6.STYLE_NAMES["body"])
    apply_style_rule(body_style, document, SECTION6.THESIS_STYLE_RULES[SECTION6.STYLE_NAMES["body"]])
    for style_name, rule in SECTION6.CUSTOM_STYLE_RULES.items():
        style = get_or_add_paragraph_style(document, style_name)
        apply_style_rule(style, document, rule)

    add_demo_content(document)
    document.save(output_path)
    print(f"[6.1] 已生成自定义样式示例文档: {output_path}")


if __name__ == "__main__":
    main()
