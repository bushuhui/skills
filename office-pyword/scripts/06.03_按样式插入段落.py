from __future__ import annotations

from pathlib import Path

from docx import Document

from _06_docx_styles import prepare_common_styles, set_default_style
from _10_style_presets import STYLE_NAMES

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = TESTS_DIR / "6.3_带样式段落插入示例.docx"

    document = Document()
    set_default_style(document)
    prepare_common_styles(document)

    document.add_paragraph("样式驱动的章节内容示例", style=STYLE_NAMES["custom_lead"])
    document.add_paragraph("绪论", style=STYLE_NAMES["title1"])
    document.add_paragraph("研究背景", style=STYLE_NAMES["title2"])
    document.add_paragraph("行业需求", style=STYLE_NAMES["title3"])
    document.add_paragraph(
        "本段正文在插入时直接指定 ThesisBody 样式，因此字体、行距、两端对齐和首行缩进均由样式统一控制。",
        style=STYLE_NAMES["body"],
    )
    document.add_paragraph(
        "引用说明段落使用 CustomQuote 样式，适合放置概述、摘录或需要区别于正文的内容。",
        style=STYLE_NAMES["custom_quote"],
    )
    document.add_paragraph(
        "注：样式化插入的重点是让段落在创建时就带上规范格式，而不是事后逐段补直接格式。",
        style=STYLE_NAMES["custom_note"],
    )

    document.save(output_path)
    print(f"[6.3] 已生成带样式段落插入示例文档: {output_path}")


if __name__ == "__main__":
    main()
