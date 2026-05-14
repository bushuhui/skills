from __future__ import annotations

from pathlib import Path

from docx import Document

from _06_docx_styles import get_style, prepare_common_styles, set_default_style
from _10_style_presets import AUTO_FORMAT_PATTERNS, AUTO_FORMAT_SAMPLE_LINES, STYLE_NAMES

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def create_source_document(source_path: Path) -> None:
    document = Document()
    for line in AUTO_FORMAT_SAMPLE_LINES:
        document.add_paragraph(line)
    document.save(source_path)


def infer_style_name(text: str) -> str:
    stripped = text.strip()
    for pattern, style_name in AUTO_FORMAT_PATTERNS:
        if pattern.match(stripped):
            return style_name
    return STYLE_NAMES["body"]


def format_existing_document(source_path: Path, output_path: Path) -> None:
    document = Document(source_path)
    set_default_style(document)
    prepare_common_styles(document)
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            paragraph.style = get_style(document, infer_style_name(paragraph.text))
    document.save(output_path)


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    source_path = TESTS_DIR / "6.4_自动格式化源文档.docx"
    output_path = TESTS_DIR / "6.4_自动格式化示例.docx"

    create_source_document(source_path)
    format_existing_document(source_path, output_path)

    print(f"[6.4] 已生成自动格式化源文档: {source_path}")
    print(f"[6.4] 已生成自动格式化示例文档: {output_path}")


if __name__ == "__main__":
    main()
