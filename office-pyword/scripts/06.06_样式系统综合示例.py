from __future__ import annotations

from pathlib import Path

from docx import Document

from _05_docx_numbering import apply_numbering, ensure_numbering
from _06_docx_styles import prepare_common_styles, set_default_style
from _10_style_presets import NUMBERING_LEVELS, STYLE_NAMES

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def build_numbering(document: Document) -> int:
    return ensure_numbering(
        document,
        [
            {
                "ilvl": level_rule["ilvl"],
                "num_fmt": level_rule["num_fmt"],
                "lvl_text": level_rule["lvl_text"],
                "suff": level_rule["suff"],
                "p_style_id": document.styles[level_rule["style_name"]].style_id,
                "left": 0,
                "hanging": 0,
            }
            for level_rule in NUMBERING_LEVELS
        ],
    )


def add_heading(document: Document, text: str, level: int, num_id: int) -> None:
    paragraph = document.add_paragraph(text, style=STYLE_NAMES[f"title{level}"])
    apply_numbering(paragraph, num_id, level - 1)


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = TESTS_DIR / "6.6_样式系统综合示例.docx"

    document = Document()
    set_default_style(document)
    prepare_common_styles(document)
    num_id = build_numbering(document)

    document.add_paragraph("样式系统与自动格式化综合示例", style=STYLE_NAMES["custom_lead"])
    document.add_paragraph(
        "说明：本页同时演示自定义样式、标题样式、正文样式、带样式插入和多级自动编号。",
        style=STYLE_NAMES["custom_note"],
    )

    add_heading(document, "绪论", 1, num_id)
    document.add_paragraph("样式系统的核心目标是把标题、正文和说明文字的格式规则集中管理，减少散落在各个脚本中的直接格式设置。", style=STYLE_NAMES["body"])
    add_heading(document, "研究背景", 2, num_id)
    document.add_paragraph("在批量生成研究报告、项目文档和技术总结时，样式驱动比逐段直接格式更稳定，也更利于后续统一修改。", style=STYLE_NAMES["body"])
    add_heading(document, "样式驱动排版", 3, num_id)
    document.add_paragraph("当标题和正文都绑定到固定样式后，插入段落时只需要指定样式名称，而不必反复设定字体、缩进、段前段后和行距。", style=STYLE_NAMES["body"])
    add_heading(document, "自动编号机制", 2, num_id)
    document.add_paragraph("本示例把编号信息写进 numbering.xml，再通过样式和段落 numPr 关联到多级标题，从而避免手工拼接编号文本。", style=STYLE_NAMES["body"])
    document.add_paragraph(
        "扩展建议：下一步可以把样式系统继续接到目录生成、标题识别和格式审查流程里，形成完整的论文排版流水线。",
        style=STYLE_NAMES["custom_quote"],
    )

    document.save(output_path)
    print(f"[6.6] 已生成样式系统综合示例文档: {output_path}")


if __name__ == "__main__":
    main()
