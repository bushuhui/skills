from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"
SOURCE_DOCX = TESTS_DIR / "测试文档.docx"
OUTPUT_DOCX = TESTS_DIR / "8.1_分页符示例.docx"


def copy_docx(source: Path, target: Path) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(source, target)


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def find_target_paragraph(document: Document):
    for paragraph in document.paragraphs:
        if normalize_text(paragraph.text) == "参考文献":
            return paragraph
    if len(document.paragraphs) > 200:
        return document.paragraphs[200]
    return document.paragraphs[-1]


def insert_page_break_before(paragraph) -> None:
    break_paragraph = paragraph.insert_paragraph_before()
    break_paragraph.add_run().add_break(WD_BREAK.PAGE)


def main() -> None:
    copy_docx(SOURCE_DOCX, OUTPUT_DOCX)
    document = Document(OUTPUT_DOCX)
    target = find_target_paragraph(document)
    insert_page_break_before(target)
    document.save(OUTPUT_DOCX)
    print(f"[8.1] 已在目标段落前插入分页符: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
