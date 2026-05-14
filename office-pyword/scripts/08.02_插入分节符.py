from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START

from _06_docx_styles import prepare_common_styles


TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"
SOURCE_DOCX = TESTS_DIR / "测试文档.docx"
OUTPUT_DOCX = TESTS_DIR / "8.2_分节符示例.docx"


def copy_docx(source: Path, target: Path) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(source, target)


def main() -> None:
    copy_docx(SOURCE_DOCX, OUTPUT_DOCX)
    document = Document(OUTPUT_DOCX)
    prepare_common_styles(document)

    document.add_paragraph("以下内容之前已插入连续分节符。", style="ThesisBody")
    continuous_section = document.add_section(WD_SECTION_START.CONTINUOUS)
    continuous_section.header.is_linked_to_previous = False
    continuous_section.footer.is_linked_to_previous = False
    continuous_section.header.paragraphs[0].text = "连续分节页眉"
    continuous_section.footer.paragraphs[0].text = "连续分节页脚"
    document.add_paragraph("这是连续分节符后的内容，可单独设置页眉页脚与页面布局。", style="ThesisBody")

    document.add_paragraph("以下内容之前已插入下一页分节符。", style="ThesisBody")
    next_page_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    next_page_section.header.is_linked_to_previous = False
    next_page_section.footer.is_linked_to_previous = False
    next_page_section.header.paragraphs[0].text = "下一页分节页眉"
    next_page_section.footer.paragraphs[0].text = "下一页分节页脚"
    document.add_paragraph("这是下一页分节符后的内容。", style="ThesisBody")

    document.save(OUTPUT_DOCX)
    print(f"[8.2] 已生成分节符示例文档: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()

