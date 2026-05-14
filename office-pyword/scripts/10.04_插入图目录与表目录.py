from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from _02_docx_captions import add_caption_paragraph
from _03_docx_fields import add_catalog_field, enable_update_fields_on_open
from _06_docx_styles import prepare_common_styles
from _07_docx_tables import apply_three_line_table, format_cell_text
from _11_docx_enhance_presets import SAMPLE_IMAGE_1, SAMPLE_TABLE_DATA

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def build_table(document):
    table = document.add_table(rows=len(SAMPLE_TABLE_DATA), cols=len(SAMPLE_TABLE_DATA[0]))
    for row_idx, row_data in enumerate(SAMPLE_TABLE_DATA):
        for col_idx, value in enumerate(row_data):
            table.cell(row_idx, col_idx).text = value
            format_cell_text(table.cell(row_idx, col_idx), bold=(row_idx == 0))
    apply_three_line_table(table)
    return table


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = TESTS_DIR / "10.4_图目录表目录示例.docx"

    document = Document()
    prepare_common_styles(document)
    enable_update_fields_on_open(document)

    document.add_paragraph("图目录", style="ToCTitleCn")
    fig_catalog = document.add_paragraph()
    add_catalog_field(fig_catalog, "图")

    document.add_paragraph("表目录", style="ToCTitleCn")
    table_catalog = document.add_paragraph()
    add_catalog_field(table_catalog, "表")
    document.add_page_break()

    document.add_paragraph("插图与表格示例", style="ThesisTitle1")
    document.add_paragraph("下面的内容用于演示图题、表题以及图目录、表目录的自动生成路径。", style="ThesisBody")

    document.add_paragraph("")
    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(SAMPLE_IMAGE_1), width=Cm(8))
    add_caption_paragraph(document, "图", "图", "三维重建结果示意图", "FigureCaption", "fig_result_demo")

    document.add_paragraph("")
    add_caption_paragraph(document, "表", "表", "实验指标对比表", "TableCaption", "tbl_metric_demo")
    build_table(document)

    document.save(output_path)
    print(f"[10.4] 已生成图目录和表目录示例文档: {output_path}")


if __name__ == "__main__":
    main()

