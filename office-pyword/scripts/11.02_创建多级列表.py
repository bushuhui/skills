from __future__ import annotations

from pathlib import Path

from docx import Document

from _05_docx_numbering import apply_numbering, ensure_numbering
from _06_docx_styles import prepare_common_styles

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = TESTS_DIR / "11.2_多级列表示例.docx"

    document = Document()
    prepare_common_styles(document)

    multi_num_id = ensure_numbering(
        document,
        [
            {"ilvl": 0, "num_fmt": "decimal", "lvl_text": "%1.", "left": 360, "hanging": 360},
            {"ilvl": 1, "num_fmt": "decimal", "lvl_text": "%1.%2.", "left": 720, "hanging": 360},
            {"ilvl": 2, "num_fmt": "decimal", "lvl_text": "%1.%2.%3.", "left": 1080, "hanging": 360},
        ],
    )

    document.add_paragraph("多级列表示例", style="ThesisTitle1")
    entries = [
        ("总体流程", 0),
        ("数据采集", 1),
        ("多视角图像", 2),
        ("点云处理", 2),
        ("模型构建", 1),
        ("结果分析", 0),
    ]
    for text, level in entries:
        paragraph = document.add_paragraph(text, style="ListBody")
        apply_numbering(paragraph, multi_num_id, level)

    document.save(output_path)
    print(f"[11.2] 已生成多级列表示例文档: {output_path}")


if __name__ == "__main__":
    main()

