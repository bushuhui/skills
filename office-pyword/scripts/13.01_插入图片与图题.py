from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from _02_docx_captions import add_caption_paragraph
from _06_docx_styles import prepare_common_styles
from _11_docx_enhance_presets import SAMPLE_IMAGE_1, SAMPLE_IMAGE_2

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def add_image_block(document, image_path: Path, width_cm: float, caption_text: str, bookmark_name: str):
    document.add_paragraph("")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm))
    add_caption_paragraph(document, "图", "图", caption_text, "FigureCaption", bookmark_name)


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = TESTS_DIR / "13.1_图片插入和图题示例.docx"

    document = Document()
    prepare_common_styles(document)

    document.add_paragraph("图片插入示例", style="ThesisTitle1")
    document.add_paragraph("本示例演示嵌入式图片插入、按比例缩放、单倍行距图片段落以及图题设置。", style="ThesisBody")

    add_image_block(document, SAMPLE_IMAGE_1, 8, "三维重建结果一", "fig_insert_01")
    add_image_block(document, SAMPLE_IMAGE_2, 6, "三维重建结果二", "fig_insert_02")

    document.save(output_path)
    print(f"[13.1] 已生成图片插入和图题示例文档: {output_path}")


if __name__ == "__main__":
    main()

