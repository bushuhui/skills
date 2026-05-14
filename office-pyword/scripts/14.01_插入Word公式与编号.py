from __future__ import annotations

from pathlib import Path

from docx import Document

from _01_docx_bookmarks import add_bookmark_to_paragraph
from _03_docx_fields import add_seq_field
from _04_docx_math import append_equation, math_fraction, math_superscript
from _06_docx_styles import prepare_common_styles
from _07_docx_tables import hide_table_borders

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def add_equation_block(document, bookmark_name: str):
    document.add_paragraph("")
    table = document.add_table(rows=1, cols=3)
    hide_table_borders(table)

    left_cell, center_cell, right_cell = table.rows[0].cells
    left_cell.text = ""

    center_paragraph = center_cell.paragraphs[0]
    append_equation(center_paragraph, ["E = m", math_superscript("c", "2"), " + ", math_fraction("1", "2"), "m", math_superscript("v", "2")])

    right_paragraph = right_cell.paragraphs[0]
    right_paragraph.style = document.styles["EquationCaption"]
    right_paragraph.add_run("（")
    add_seq_field(right_paragraph, "式", "1")
    right_paragraph.add_run("）")
    add_bookmark_to_paragraph(document, right_paragraph, bookmark_name)
    document.add_paragraph("")


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = TESTS_DIR / "14.1_Word公式与编号示例.docx"

    document = Document()
    prepare_common_styles(document)

    document.add_paragraph("Word 原生公式与编号示例", style="ThesisTitle1")
    document.add_paragraph("本示例插入 Word 原生 OMML 公式，并采用无边框表格实现公式居中、编号右对齐。", style="ThesisBody")
    add_equation_block(document, "eq_energy_demo")
    document.add_paragraph("公式块与上下文之间通过额外的半行空白段落进行控制。", style="ThesisBody")

    document.save(output_path)
    print(f"[14.1] 已生成 Word 原生公式与编号示例文档: {output_path}")


if __name__ == "__main__":
    main()

