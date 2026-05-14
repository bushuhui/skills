from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

from _04_docx_math import append_equation, math_fraction
from _06_docx_styles import prepare_common_styles

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def build_source_formula_document(source_path: Path) -> None:
    document = Document()
    prepare_common_styles(document)
    document.add_paragraph("外部公式源文档", style="ThesisTitle1")
    paragraph = document.add_paragraph()
    append_equation(paragraph, ["x = ", math_fraction("a+b", "c+d")])
    document.save(source_path)


def import_first_formula_paragraph(source_path: Path, target_document: Document) -> None:
    source_document = Document(source_path)
    source_paragraph = source_document.paragraphs[1]
    target_paragraph = target_document.add_paragraph()
    for child in source_paragraph._p:
        target_paragraph._p.append(deepcopy(child))


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    source_path = TESTS_DIR / "14.2_公式源文档.docx"
    output_path = TESTS_DIR / "14.2_MathType兼容公式示例.docx"

    build_source_formula_document(source_path)

    document = Document()
    prepare_common_styles(document)
    document.add_paragraph("MathType 兼容公式插入示例", style="ThesisTitle1")
    document.add_paragraph("本示例演示从外部公式文档导入公式段落 XML 的兼容路径，适合后续扩展到 MathType/外部公式来源。", style="ThesisBody")
    import_first_formula_paragraph(source_path, document)

    document.save(output_path)
    print(f"[14.2] 已生成公式源文档: {source_path}")
    print(f"[14.2] 已生成 MathType 兼容公式示例文档: {output_path}")


if __name__ == "__main__":
    main()

