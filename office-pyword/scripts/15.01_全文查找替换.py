from __future__ import annotations

from pathlib import Path

from docx import Document

from _06_docx_styles import prepare_common_styles
from _08_docx_textops import replace_text_everywhere

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def build_source_document(source_path: Path) -> None:
    document = Document()
    prepare_common_styles(document)
    section = document.sections[0]
    section.header.paragraphs[0].text = "项目名称：digital twin"
    section.footer.paragraphs[0].text = "图1 digital twin"

    document.add_paragraph("全文查找与替换示例", style="ThesisTitle1")
    document.add_paragraph("digital twin 技术正在重塑空间计算流程。", style="ThesisBody")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "值"
    table.cell(1, 0).text = "关键词"
    table.cell(1, 1).text = "digital twin"
    document.save(source_path)


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    source_path = TESTS_DIR / "15.1_查找替换源文档.docx"
    output_path = TESTS_DIR / "15.1_查找替换结果.docx"

    build_source_document(source_path)
    document = Document(source_path)
    replacements = [("digital twin", "digital twin system")]
    count = replace_text_everywhere(document, replacements)
    document.save(output_path)

    print(f"[15.1] 已生成查找替换源文档: {source_path}")
    print(f"[15.1] 已生成查找替换结果文档: {output_path}")
    print(f"[15.1] 替换次数: {count}")


if __name__ == "__main__":
    main()

