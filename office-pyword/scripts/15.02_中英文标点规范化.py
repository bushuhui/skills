from __future__ import annotations

from pathlib import Path

from docx import Document

from _06_docx_styles import prepare_common_styles
from _08_docx_textops import replace_text_everywhere
from _11_docx_enhance_presets import PUNCTUATION_REPLACEMENTS_ZH

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = TESTS_DIR / "15.2_中英文标点规范化示例.docx"

    document = Document()
    prepare_common_styles(document)
    document.add_paragraph("中英文标点规范化示例", style="ThesisTitle1")
    document.add_paragraph("研究目标: 提升模型精度, 提高效率; 并完善输出格式 (适用于中文报告).", style="ThesisBody")

    replace_text_everywhere(document, PUNCTUATION_REPLACEMENTS_ZH)
    document.save(output_path)

    print(f"[15.2] 已生成中英文标点规范化示例文档: {output_path}")


if __name__ == "__main__":
    main()

