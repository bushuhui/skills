from __future__ import annotations

from pathlib import Path

from docx import Document

from _06_docx_styles import get_style, prepare_common_styles
from _08_docx_textops import replace_text_everywhere
from _11_docx_enhance_presets import REVIEW_FIXES

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def build_source_document(source_path: Path) -> None:
    document = Document()
    prepare_common_styles(document)
    document.add_paragraph("绪论")
    document.add_paragraph("本节介绍 digital twin 的应用。", style="ThesisBody")
    document.add_paragraph("图1 为系统流程图。", style="ThesisBody")
    document.save(source_path)


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    source_path = TESTS_DIR / "15.3_自动修正源文档.docx"
    output_path = TESTS_DIR / "15.3_自动修正结果.docx"

    build_source_document(source_path)
    document = Document(source_path)
    prepare_common_styles(document)

    text_replacements = [(item["find"], item["replace"]) for item in REVIEW_FIXES if item["action"] == "replace_text"]
    replace_text_everywhere(document, text_replacements)

    for item in REVIEW_FIXES:
        if item["action"] == "apply_style_by_text":
            for paragraph in document.paragraphs:
                if item["contains"] in paragraph.text:
                    paragraph.style = get_style(document, item["style_name"])

    document.save(output_path)
    print(f"[15.3] 已生成自动修正源文档: {source_path}")
    print(f"[15.3] 已生成自动修正结果文档: {output_path}")


if __name__ == "__main__":
    main()

