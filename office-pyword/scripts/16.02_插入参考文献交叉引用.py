from __future__ import annotations

from pathlib import Path

from docx import Document

from _01_docx_bookmarks import add_bookmark_around_run
from _03_docx_fields import add_ref_field, enable_update_fields_on_open
from _06_docx_styles import prepare_common_styles
from _11_docx_enhance_presets import SAMPLE_BIBLIOGRAPHY

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = TESTS_DIR / "16.2_参考文献交叉引用示例.docx"

    document = Document()
    prepare_common_styles(document)
    enable_update_fields_on_open(document)

    document.add_paragraph("正文中的参考文献交叉引用示例", style="ThesisTitle1")
    paragraph = document.add_paragraph(style="ThesisBody")
    paragraph.add_run("相关研究可参见 ")
    add_ref_field(paragraph, SAMPLE_BIBLIOGRAPHY[0]["bookmark"], "[1]")
    paragraph.add_run(" 和 ")
    add_ref_field(paragraph, SAMPLE_BIBLIOGRAPHY[1]["bookmark"], "[2]")
    paragraph.add_run("。")

    document.add_paragraph("参考文献", style="ThesisTitle1")
    for entry in SAMPLE_BIBLIOGRAPHY:
        ref_paragraph = document.add_paragraph(style="BibliographyEntry")
        number_run = ref_paragraph.add_run(f"[{entry['number']}] ")
        add_bookmark_around_run(document, number_run, entry["bookmark"])
        ref_paragraph.add_run(entry["text"])

    document.save(output_path)
    print(f"[16.2] 已生成参考文献交叉引用示例文档: {output_path}")


if __name__ == "__main__":
    main()

