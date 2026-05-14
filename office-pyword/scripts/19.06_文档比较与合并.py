from __future__ import annotations

import json
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document

from _06_docx_styles import prepare_common_styles
from _14_docx_review_demo import TESTS_DIR
from _16_docx_revisions import (
    build_package_map,
    ensure_track_revisions_setting,
    get_top_level_paragraphs,
    next_revision_id,
    normalize_text,
    paragraph_diff_segments,
    parse_document_root,
    process_revisions_docx,
    replace_paragraph_segments,
    update_document_root,
    write_package_map,
)

LEFT_DOCX = TESTS_DIR / "19.6_比较源文档_A.docx"
RIGHT_DOCX = TESTS_DIR / "19.6_比较源文档_B.docx"
MERGED_DOCX = TESTS_DIR / "19.6_比较合并结果.docx"
ACCEPTED_DOCX = TESTS_DIR / "19.6_合并后接受修订结果.docx"
REJECTED_DOCX = TESTS_DIR / "19.6_合并后拒绝修订结果.docx"
REPORT_JSON = TESTS_DIR / "19.6_比较报告.json"

AUTHOR = "PyWord Compare"


def build_demo_document(path: Path, paragraphs: list[tuple[str, str]]) -> None:
    document = Document()
    prepare_common_styles(document)
    for style_name, text in paragraphs:
        document.add_paragraph(text, style=style_name)
    document.save(path)


def collect_paragraph_records(path: Path) -> list[dict[str, str]]:
    document = Document(path)
    rows: list[dict[str, str]] = []
    for paragraph in document.paragraphs:
        if not normalize_text(paragraph.text):
            continue
        rows.append(
            {
                "text": paragraph.text,
                "style_name": paragraph.style.name if paragraph.style is not None else "Normal",
            }
        )
    return rows


def build_merge_plan(left_rows: list[dict[str, str]], right_rows: list[dict[str, str]]) -> tuple[list[dict[str, object]], list[dict[str, object]]]:
    left_keys = [normalize_text(item["text"]) for item in left_rows]
    right_keys = [normalize_text(item["text"]) for item in right_rows]
    matcher = SequenceMatcher(a=left_keys, b=right_keys)

    plan: list[dict[str, object]] = []
    diff_rows: list[dict[str, object]] = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for row in right_rows[j1:j2]:
                plan.append({"style_name": row["style_name"], "seed_text": row["text"], "segments": None})
            continue

        left_block = left_rows[i1:i2]
        right_block = right_rows[j1:j2]

        if tag == "replace":
            paired_count = min(len(left_block), len(right_block))
            for index in range(paired_count):
                left_row = left_block[index]
                right_row = right_block[index]
                segments = paragraph_diff_segments(left_row["text"], right_row["text"])
                plan.append(
                    {
                        "style_name": right_row["style_name"] or left_row["style_name"],
                        "seed_text": right_row["text"],
                        "segments": segments,
                    }
                )
                diff_rows.append(
                    {
                        "type": "replace",
                        "from": left_row["text"],
                        "to": right_row["text"],
                        "segments": segments,
                    }
                )
            for row in left_block[paired_count:]:
                plan.append({"style_name": row["style_name"], "seed_text": row["text"], "segments": [("delete", row["text"])]})
                diff_rows.append({"type": "delete", "from": row["text"], "to": ""})
            for row in right_block[paired_count:]:
                plan.append({"style_name": row["style_name"], "seed_text": row["text"], "segments": [("insert", row["text"])]})
                diff_rows.append({"type": "insert", "from": "", "to": row["text"]})
            continue

        if tag == "delete":
            for row in left_block:
                plan.append({"style_name": row["style_name"], "seed_text": row["text"], "segments": [("delete", row["text"])]})
                diff_rows.append({"type": "delete", "from": row["text"], "to": ""})
            continue

        if tag == "insert":
            for row in right_block:
                plan.append({"style_name": row["style_name"], "seed_text": row["text"], "segments": [("insert", row["text"])]})
                diff_rows.append({"type": "insert", "from": "", "to": row["text"]})

    return plan, diff_rows


def build_merged_document(output_path: Path, plan: list[dict[str, object]]) -> None:
    document = Document()
    prepare_common_styles(document)
    available_styles = {style.name for style in document.styles}

    for row in plan:
        style_name = row["style_name"] if row["style_name"] in available_styles else "ThesisBody"
        document.add_paragraph(str(row["seed_text"]), style=style_name)
    document.save(output_path)

    package_map = build_package_map(output_path)
    ensure_track_revisions_setting(package_map, enabled=True)
    document_root = parse_document_root(package_map)
    top_level_paragraphs = get_top_level_paragraphs(document_root)
    timestamp_id = next_revision_id(document_root)

    for paragraph_element, row in zip(top_level_paragraphs, plan):
        segments = row["segments"]
        if not segments:
            continue
        timestamp_id = replace_paragraph_segments(
            paragraph_element,
            segments,
            AUTHOR,
            timestamp_id,
        )

    update_document_root(package_map, document_root)
    write_package_map(package_map, output_path)


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)

    left_paragraphs = [
        ("ThesisTitle1", "文档比较与合并示例"),
        ("ThesisBody", "本报告围绕三维重建数据处理流程展开说明。"),
        ("ThesisTitle2", "研究背景"),
        ("ThesisBody", "多视角影像采集能够为三维重建提供稳定的数据基础。"),
        ("ThesisBody", "点云配准与网格重建是实验流程中的关键环节。"),
        ("ThesisBody", "传统手工配准流程仍然具有一定参考价值。"),
        ("ThesisTitle2", "实验结论"),
        ("ThesisBody", "实验结果表明该流程具有良好的精度与稳定性。"),
    ]
    right_paragraphs = [
        ("ThesisTitle1", "文档比较与合并示例"),
        ("ThesisBody", "本报告围绕三维重建数据处理与质量控制流程展开说明。"),
        ("ThesisTitle2", "研究背景"),
        ("ThesisBody", "多视角影像采集能够为三维重建提供更高分辨率且稳定的数据基础。"),
        ("ThesisBody", "点云配准与网格重建是实验流程中的核心环节。"),
        ("ThesisBody", "自动化质检步骤进一步提升了结果一致性。"),
        ("ThesisTitle2", "实验结论"),
        ("ThesisBody", "实验结果表明该流程在精度、稳定性与自动化程度方面均有明显提升。"),
    ]

    build_demo_document(LEFT_DOCX, left_paragraphs)
    build_demo_document(RIGHT_DOCX, right_paragraphs)

    left_rows = collect_paragraph_records(LEFT_DOCX)
    right_rows = collect_paragraph_records(RIGHT_DOCX)
    plan, diff_rows = build_merge_plan(left_rows, right_rows)

    build_merged_document(MERGED_DOCX, plan)
    accepted_counts = process_revisions_docx(MERGED_DOCX, ACCEPTED_DOCX, accept=True)
    rejected_counts = process_revisions_docx(MERGED_DOCX, REJECTED_DOCX, accept=False)

    report = {
        "left_document": str(LEFT_DOCX),
        "right_document": str(RIGHT_DOCX),
        "merged_document": str(MERGED_DOCX),
        "accepted_document": str(ACCEPTED_DOCX),
        "rejected_document": str(REJECTED_DOCX),
        "diff_count": len(diff_rows),
        "diff_rows": diff_rows,
        "accepted_summary": accepted_counts,
        "rejected_summary": rejected_counts,
    }
    REPORT_JSON.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"[19.6] 已生成比较源文档 A: {LEFT_DOCX}")
    print(f"[19.6] 已生成比较源文档 B: {RIGHT_DOCX}")
    print(f"[19.6] 已生成带修订的比较合并文档: {MERGED_DOCX}")
    print(f"[19.6] 已生成接受修订后的合并结果: {ACCEPTED_DOCX}")
    print(f"[19.6] 已生成拒绝修订后的合并结果: {REJECTED_DOCX}")
    print(f"[19.6] 已输出比较报告: {REPORT_JSON}")


if __name__ == "__main__":
    main()
