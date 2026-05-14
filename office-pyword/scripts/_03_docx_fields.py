from __future__ import annotations

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


def _add_run_with_fld_char(paragraph, fld_char_type: str, dirty: bool = False):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), fld_char_type)
    if dirty:
        fld_char.set(qn("w:dirty"), "true")
    run._r.append(fld_char)
    return run


def _add_instr_text_run(paragraph, instruction: str):
    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(XML_SPACE, "preserve")
    instr.text = instruction
    run._r.append(instr)
    return run


def add_complex_field(paragraph, instruction: str, placeholder_text: str = ""):
    _add_run_with_fld_char(paragraph, "begin", dirty=True)
    _add_instr_text_run(paragraph, instruction)
    _add_run_with_fld_char(paragraph, "separate")
    if placeholder_text:
        paragraph.add_run(placeholder_text)
    _add_run_with_fld_char(paragraph, "end")


def add_toc_field(paragraph, levels: str = "1-3") -> None:
    add_complex_field(paragraph, f'TOC \\o "{levels}" \\h \\z \\u', "目录将在 Word 中更新")


def add_catalog_field(paragraph, label: str) -> None:
    add_complex_field(paragraph, f'TOC \\h \\z \\c "{label}"', f"{label}目录将在 Word 中更新")


def add_seq_field(paragraph, identifier: str, placeholder_text: str = "1") -> None:
    add_complex_field(paragraph, f"SEQ {identifier} \\* ARABIC", placeholder_text)


def add_ref_field(paragraph, bookmark_name: str, placeholder_text: str = "") -> None:
    add_complex_field(paragraph, f" REF {bookmark_name} \\h ", placeholder_text or bookmark_name)


def add_pageref_field(paragraph, bookmark_name: str, placeholder_text: str = "1") -> None:
    add_complex_field(paragraph, f" PAGEREF {bookmark_name} \\h ", placeholder_text)


def enable_update_fields_on_open(document) -> None:
    settings = document.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
