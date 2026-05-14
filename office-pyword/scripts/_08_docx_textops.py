from __future__ import annotations

from copy import deepcopy

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


def iter_all_paragraphs(document):
    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for paragraph in section.footer.paragraphs:
            yield paragraph


def clone_run_with_text(run, text: str):
    new_r = deepcopy(run._element)
    for child in list(new_r):
        if child.tag != qn("w:rPr"):
            new_r.remove(child)
    text_element = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        text_element.set(XML_SPACE, "preserve")
    text_element.text = text
    new_r.append(text_element)
    return new_r


def build_run_spans(paragraph):
    spans = []
    cursor = 0
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        spans.append({"run": run, "start": cursor, "end": cursor + len(text)})
        cursor += len(text)
    return spans


def clone_original_range(spans, start: int, end: int):
    if start >= end:
        return []
    elements = []
    for item in spans:
        overlap_start = max(start, item["start"])
        overlap_end = min(end, item["end"])
        if overlap_start >= overlap_end:
            continue
        left = overlap_start - item["start"]
        right = overlap_end - item["start"]
        elements.append(clone_run_with_text(item["run"], item["run"].text[left:right]))
    return elements


def find_matches(text: str, needle: str):
    matches = []
    start = 0
    while True:
        index = text.find(needle, start)
        if index == -1:
            break
        matches.append((index, index + len(needle)))
        start = index + len(needle)
    return matches


def replace_text_in_paragraph(paragraph, needle: str, replacement: str) -> int:
    full_text = paragraph.text
    matches = find_matches(full_text, needle)
    if not matches:
        return 0

    spans = build_run_spans(paragraph)
    if not spans:
        return 0

    new_elements = []
    cursor = 0
    for match_start, match_end in matches:
        new_elements.extend(clone_original_range(spans, cursor, match_start))
        template = next(item["run"] for item in spans if item["start"] < match_end and item["end"] > match_start)
        new_elements.append(clone_run_with_text(template, replacement))
        cursor = match_end
    new_elements.extend(clone_original_range(spans, cursor, len(full_text)))

    original_runs = list(paragraph.runs)
    for run in original_runs:
        paragraph._p.remove(run._element)
    for element in new_elements:
        paragraph._p.append(element)
    return len(matches)


def replace_text_everywhere(document, replacements: list[tuple[str, str]]) -> int:
    replaced_count = 0
    for paragraph in iter_all_paragraphs(document):
        for find_text, replace_text in replacements:
            replaced_count += replace_text_in_paragraph(paragraph, find_text, replace_text)
    return replaced_count
