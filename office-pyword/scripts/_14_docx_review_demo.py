from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Cm

from _06_docx_styles import prepare_common_styles

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"
SOURCE_DOCX = TESTS_DIR / "19_示例源文档.docx"


def ensure_demo_source_document(path: Path | None = None) -> Path:
    target = path or SOURCE_DOCX
    if target.exists():
        return target

    target.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    prepare_common_styles(document)

    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(2.54)

    document.add_paragraph("文档安全与审阅示例", style="ThesisTitle1")
    document.add_paragraph("本示例文档用于演示限制编辑密码保护、审阅意见写入与评论提取。", style="ThesisBody")
    document.add_paragraph("研究背景", style="ThesisTitle2")
    document.add_paragraph("多视角影像采集能够为三维重建提供稳定的数据基础。", style="ThesisBody")
    document.add_paragraph("点云配准与网格重建是实验流程中的关键环节。", style="ThesisBody")
    document.add_paragraph("实验结论", style="ThesisTitle2")
    document.add_paragraph("建议在正式论文中补充误差分析与参考文献对比。", style="ThesisBody")

    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "指标"
    table.cell(0, 1).text = "结果"
    table.cell(1, 0).text = "重投影误差"
    table.cell(1, 1).text = "0.42 px"
    table.cell(2, 0).text = "点云密度"
    table.cell(2, 1).text = "18.6 pts/cm2"

    document.save(target)
    return target
