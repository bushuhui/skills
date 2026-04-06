from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"

JIANG_JIN_JIU_LINES = [
    "将进酒",
    "李白",
    "",
    "君不见，黄河之水天上来，奔流到海不复回。",
    "君不见，高堂明镜悲白发，朝如青丝暮成雪。",
    "",
    "人生得意须尽欢，莫使金樽空对月。",
    "天生我材必有用，千金散尽还复来。",
    "",
    "烹羊宰牛且为乐，会须一饮三百杯。",
    "岑夫子，丹丘生，将进酒，杯莫停。",
    "",
    "与君歌一曲，请君为我倾耳听。",
    "钟鼓馔玉不足贵，但愿长醉不愿醒。",
    "古来圣贤皆寂寞，惟有饮者留其名。",
    "",
    "陈王昔时宴平乐，斗酒十千恣欢谑。",
    "主人何为言少钱，径须沽取对君酌。",
    "",
    "五花马，千金裘，呼儿将出换美酒，",
    "与尔同销万古愁。",
]


def ensure_tests_dir() -> Path:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    return TESTS_DIR


def set_default_style(document, font_name: str = "宋体", font_size_pt: int = 12) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = font_name
    normal_style.font.size = Pt(font_size_pt)

    style_rpr = normal_style.element.get_or_add_rPr()
    r_fonts = style_rpr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        style_rpr.append(r_fonts)

    for key in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        r_fonts.set(qn(key), font_name)

    normal_format = normal_style.paragraph_format
    normal_format.line_spacing = 1.5
    normal_format.space_before = Pt(0)
    normal_format.space_after = Pt(0)


def apply_run_style(run, font_name: str = "宋体", font_size_pt: int = 12, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.bold = bold

    run_rpr = run._element.get_or_add_rPr()
    r_fonts = run_rpr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        run_rpr.append(r_fonts)

    for key in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        r_fonts.set(qn(key), font_name)


def apply_paragraph_format(
    paragraph,
    alignment=None,
    line_spacing: float = 1.5,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
) -> None:
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_before = Pt(space_before_pt)
    paragraph.paragraph_format.space_after = Pt(space_after_pt)


def create_document():
    document = Document()
    set_default_style(document)
    return document


def main() -> None:
    output_dir = ensure_tests_dir()
    output_path = output_dir / "1.1_将进酒.docx"

    document = create_document()

    for index, line in enumerate(JIANG_JIN_JIU_LINES):
        paragraph = document.add_paragraph()
        if not line:
            continue

        run = paragraph.add_run(line)
        if index == 0:
            apply_run_style(run, font_name="黑体", font_size_pt=18, bold=True)
            apply_paragraph_format(paragraph, WD_ALIGN_PARAGRAPH.CENTER, 1.25, 10, 10)
        elif index == 1:
            apply_run_style(run, font_name="楷体", font_size_pt=14)
            apply_paragraph_format(paragraph, WD_ALIGN_PARAGRAPH.CENTER, 1.25, 0, 10)
        else:
            apply_run_style(run, font_name="宋体", font_size_pt=12)
            apply_paragraph_format(paragraph, WD_ALIGN_PARAGRAPH.CENTER, 1.5, 0, 4)

    document.save(output_path)
    print(f"[1.1] 已生成新 Word 文档: {output_path}")


if __name__ == "__main__":
    main()
