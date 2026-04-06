import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"

THREE_D_REPORT_MARKDOWN = """# 三维重建技术报告

## 摘要
三维重建技术通过多视角图像、深度传感器或激光点云恢复场景的几何形态，是数字孪生、智能制造和沉浸式内容生产的重要基础能力。
本报告围绕三维重建的核心流程、典型应用与发展趋势展开说明，重点关注基于视觉的重建方案。

## 绪论
三维重建的目标是从二维观测数据中恢复目标的三维结构。随着计算机视觉、图形学和传感器技术的持续发展，该技术已经从实验室研究逐步走向工程化落地。

### 研究背景
在机器人导航、自动驾驶、文化遗产保护等场景中，系统需要准确理解真实世界的空间结构。三维重建不仅能够提升场景感知能力，还能够为后续测量、仿真与可视化提供统一的数据基础。

### 研究意义
对复杂场景进行数字化建模，有助于降低人工建模成本，并提高模型更新效率。对于高校和企业来说，掌握三维重建流程也有助于支撑科研分析、产品设计和运维决策。

## 核心技术
完整的三维重建流程通常包括数据采集、特征提取与匹配、相机位姿估计、稠密重建、表面重建以及纹理映射等步骤。不同步骤的算法选择会直接影响模型精度与计算效率。

### 数据采集与预处理
数据采集阶段通常采用单目相机、双目相机、RGB-D 设备或激光扫描仪。预处理工作包括畸变校正、图像筛选、去噪和曝光一致性调整，为后续特征匹配提供稳定输入。

### 特征匹配与位姿估计
特征匹配负责在多张图像之间建立对应关系，常见方法包括 SIFT、ORB 和 SuperPoint。位姿估计则根据匹配结果恢复相机的外参与内参关系，是稀疏重建精度的关键环节。

### 稠密重建与纹理映射
稠密重建在稀疏点云基础上恢复更完整的几何细节，常用方法包括多视图立体和体素融合。纹理映射决定了模型的视觉质量，高质量贴图能够提升成果的可视化效果，并为数字孪生、虚拟展示等场景提供更真实的表达。

## 典型应用
在工业检测领域，三维重建可以恢复零部件的真实尺寸与表面缺陷，帮助企业开展质量评估与逆向工程。在文博保护领域，三维重建能够对珍贵文物进行高精度留存，为展示、修复和学术研究提供数据支撑。

## 挑战与发展趋势
当前三维重建仍面临纹理缺失、弱纹理区域匹配困难、大规模场景计算开销高等问题。未来技术将与神经辐射场、生成式模型和多模态感知进一步结合，在实时性、鲁棒性和自动化程度方面取得更大突破。

## 结论
三维重建技术已经形成较为完整的工程链路，并在多个行业展现出显著价值。随着算法与硬件协同进步，其应用边界将继续扩展，成为数字空间构建的重要底座。
"""


def ensure_tests_dir() -> Path:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    return TESTS_DIR


def write_text_file(path: Path, text: str) -> Path:
    path.write_text(text, encoding="utf-8")
    return path


def set_default_style(document, font_name: str = "宋体", font_size_pt: int = 12) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = font_name
    normal_style.font.size = Pt(font_size_pt)

    style_rpr = normal_style.element.get_or_add_rPr()
    r_fonts = style_rpr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        style_rpr.append(r_fonts)

    for key in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        r_fonts.set(qn(key), font_name)

    normal_format = normal_style.paragraph_format
    normal_format.line_spacing = 1.5
    normal_format.space_before = Pt(0)
    normal_format.space_after = Pt(0)


def create_document():
    document = Document()
    set_default_style(document)
    return document


def apply_run_style(run, font_name: str = "宋体", font_size_pt: int = 12, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.bold = bold

    run_rpr = run._element.get_or_add_rPr()
    r_fonts = run_rpr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        run_rpr.append(r_fonts)

    for key in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        r_fonts.set(qn(key), font_name)


def apply_paragraph_format(
    paragraph,
    alignment=None,
    line_spacing: float = 1.5,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
) -> None:
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_before = Pt(space_before_pt)
    paragraph.paragraph_format.space_after = Pt(space_after_pt)


def parse_markdown_structure(markdown_text: str) -> list[dict[str, str | int]]:
    nodes: list[dict[str, str | int]] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if paragraph_lines:
            nodes.append({"type": "paragraph", "text": "".join(paragraph_lines).strip()})
            paragraph_lines.clear()

    for raw_line in markdown_text.splitlines():
        stripped = raw_line.strip()
        if not stripped:
            flush_paragraph()
            continue

        match = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if match:
            flush_paragraph()
            nodes.append(
                {
                    "type": "heading",
                    "level": len(match.group(1)),
                    "text": match.group(2).strip(),
                }
            )
            continue

        paragraph_lines.append(stripped)

    flush_paragraph()
    return nodes


def add_report_node(document, node: dict[str, str | int]) -> None:
    node_type = node["type"]
    text = str(node["text"])

    if node_type == "heading":
        level = int(node["level"])
        paragraph = document.add_paragraph(style=f"Heading {level}")
        run = paragraph.add_run(text)
        if level == 1:
            apply_run_style(run, font_name="黑体", font_size_pt=16)
            apply_paragraph_format(paragraph, WD_ALIGN_PARAGRAPH.CENTER, 1.25, 16, 16)
        elif level == 2:
            apply_run_style(run, font_name="宋体", font_size_pt=14, bold=True)
            apply_paragraph_format(paragraph, WD_ALIGN_PARAGRAPH.LEFT, 1.25, 8, 8)
        else:
            apply_run_style(run, font_name="宋体", font_size_pt=12, bold=True)
            apply_paragraph_format(paragraph, WD_ALIGN_PARAGRAPH.LEFT, 1.25, 6, 6)
        return

    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    apply_run_style(run)
    apply_paragraph_format(paragraph, WD_ALIGN_PARAGRAPH.JUSTIFY, 1.5, 0, 0)


def build_document_from_markdown(markdown_text: str, output_path: Path) -> Path:
    document = create_document()
    for node in parse_markdown_structure(markdown_text):
        add_report_node(document, node)
    document.save(output_path)
    return output_path


def main() -> None:
    output_dir = ensure_tests_dir()
    txt_path = output_dir / "1.2_三维重建技术报告.txt"
    docx_path = output_dir / "1.2_三维重建技术报告.docx"

    write_text_file(txt_path, THREE_D_REPORT_MARKDOWN)
    build_document_from_markdown(THREE_D_REPORT_MARKDOWN, docx_path)

    print(f"[1.2] 已生成结构化报告文本: {txt_path}")
    print(f"[1.2] 已根据标题节点和正文节点生成 DOCX: {docx_path}")


if __name__ == "__main__":
    main()
