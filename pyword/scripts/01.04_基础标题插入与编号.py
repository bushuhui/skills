from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def ensure_tests_dir() -> Path:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    return TESTS_DIR


def set_default_style(document, font_name: str = "宋体", font_size_pt: int = 12) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = font_name
    normal_style.font.size = Pt(font_size_pt)

    style_rpr = normal_style.element.get_or_add_rPr()
    r_fonts = style_rpr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        style_rpr.append(r_fonts)

    for key in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        r_fonts.set(qn(key), font_name)

    normal_format = normal_style.paragraph_format
    normal_format.line_spacing = 1.5
    normal_format.space_before = Pt(0)
    normal_format.space_after = Pt(0)


def apply_run_style(run, font_name: str = "宋体", font_size_pt: int = 12, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.bold = bold

    run_rpr = run._element.get_or_add_rPr()
    r_fonts = run_rpr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        run_rpr.append(r_fonts)

    for key in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        r_fonts.set(qn(key), font_name)


def apply_paragraph_format(
    paragraph,
    alignment=None,
    line_spacing: float = 1.5,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
) -> None:
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_before = Pt(space_before_pt)
    paragraph.paragraph_format.space_after = Pt(space_after_pt)


def next_heading_number(counters: list[int], level: int) -> str:
    counters[level - 1] += 1
    for index in range(level, len(counters)):
        counters[index] = 0
    return ".".join(str(value) for value in counters[:level] if value > 0)


def add_numbered_thesis_heading(document, title: str, level: int, counters: list[int]) -> str:
    number_text = next_heading_number(counters, level)
    paragraph = document.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(f"{number_text} {title}")

    if level == 1:
        apply_run_style(run, font_name="黑体", font_size_pt=16)
        apply_paragraph_format(paragraph, WD_ALIGN_PARAGRAPH.CENTER, 1.25, 16, 16)
    elif level == 2:
        apply_run_style(run, font_name="宋体", font_size_pt=14, bold=True)
        apply_paragraph_format(paragraph, WD_ALIGN_PARAGRAPH.LEFT, 1.25, 8, 8)
    else:
        apply_run_style(run, font_name="宋体", font_size_pt=12, bold=True)
        apply_paragraph_format(paragraph, WD_ALIGN_PARAGRAPH.LEFT, 1.25, 8, 8)

    return f"{number_text} {title}"


def add_body_paragraph(document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    apply_run_style(run)
    apply_paragraph_format(paragraph, WD_ALIGN_PARAGRAPH.JUSTIFY, 1.5, 0, 0)


def main() -> None:
    output_dir = ensure_tests_dir()
    output_path = output_dir / "1.4_毕业论文标题编号示例.docx"

    document = Document()
    set_default_style(document)

    counters = [0, 0, 0]
    generated_titles: list[str] = []

    generated_titles.append(add_numbered_thesis_heading(document, "绪论", 1, counters))
    add_body_paragraph(document, "本章说明三维重建技术的研究背景、研究目的与论文组织结构。")

    generated_titles.append(add_numbered_thesis_heading(document, "研究背景", 2, counters))
    add_body_paragraph(document, "随着数字孪生、机器人感知与虚拟现实的发展，三维重建已成为重要的空间信息获取手段。")

    generated_titles.append(add_numbered_thesis_heading(document, "行业发展现状", 3, counters))
    add_body_paragraph(document, "当前主流方法包括基于传统几何的 SfM/MVS 流程，以及结合 NeRF 的新型神经表示方法。")

    generated_titles.append(add_numbered_thesis_heading(document, "研究目标", 2, counters))
    add_body_paragraph(document, "研究目标是构建一套兼顾精度、效率和可扩展性的三维重建技术路线。")

    generated_titles.append(add_numbered_thesis_heading(document, "关键技术路线", 1, counters))
    add_body_paragraph(document, "本章围绕数据采集、特征匹配、位姿估计与稠密建模等关键环节展开。")

    generated_titles.append(add_numbered_thesis_heading(document, "数据采集与预处理", 2, counters))
    add_body_paragraph(document, "数据采集阶段需要保证视角覆盖、曝光稳定和标定精度，为后续重建提供高质量输入。")

    generated_titles.append(add_numbered_thesis_heading(document, "相机标定", 3, counters))
    add_body_paragraph(document, "相机标定用于估计内参和畸变参数，是提升重建尺度一致性与几何精度的重要前提。")

    generated_titles.append(add_numbered_thesis_heading(document, "实验设计", 1, counters))
    add_body_paragraph(document, "本章展示实验场景、评价指标和结果分析方式。")

    generated_titles.append(add_numbered_thesis_heading(document, "评价指标", 2, counters))
    add_body_paragraph(document, "常用评价指标包括重投影误差、点云完整率、表面精度和运行时延。")

    generated_titles.append(add_numbered_thesis_heading(document, "结果分析", 3, counters))
    add_body_paragraph(document, "结果分析既关注定量指标，也关注模型细节、纹理质量和复杂场景下的稳定性。")

    document.save(output_path)

    print(f"[1.4] 已生成毕业论文标题编号示例文档: {output_path}")
    for title in generated_titles:
        print(f"[1.4] 标题编号: {title}")


if __name__ == "__main__":
    main()
