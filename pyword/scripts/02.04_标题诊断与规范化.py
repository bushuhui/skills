import json
import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"
INPUT_DOCX = TESTS_DIR / "测试文档.docx"

FRONT_HEADINGS = {"摘要", "Abstract", "目录", "图目录", "表目录", "参考文献", "致谢", "附录"}
EXCLUDED_STYLES = {"table of figures", "图目录2", "表目录3", "表目录2", "图目录"}
STYLE_HINTS = {"Subtitle", "Title", "1正文", "heading 1", "heading 2", "heading 3"}

NUMERIC_HEADING_PATTERNS = [
    (re.compile(r"^(\d+\.\d+\.\d+)\s*(.+)$"), 3),
    (re.compile(r"^(\d+\.\d+)\s*(.+)$"), 2),
    (re.compile(r"^(\d+)\s+(.+)$"), 1),
    (re.compile(r"^(第[一二三四五六七八九十百千0-9]+章)\s*(.+)$"), 1),
    (re.compile(r"^(第[一二三四五六七八九十百千0-9]+节)\s*(.+)$"), 2),
    (re.compile(r"^([一二三四五六七八九十]+、)\s*(.+)$"), 2),
    (re.compile(r"^(（[一二三四五六七八九十]+）)\s*(.+)$"), 3),
]
CAPTION_PATTERNS = [
    re.compile(r"^[图表]\s*\d+(\.\d+)*"),
    re.compile(r"^(Fig\.?|Figure)\s*\d+", re.IGNORECASE),
    re.compile(r"^Table\s*\d+", re.IGNORECASE),
]

LEVEL_RULES = {
    1: {"font_name": "黑体", "font_size_pt": 16.0, "bold": False, "alignment": "center", "line_spacing": 1.25, "left_indent_pt": 0.0, "first_line_indent_pt": 0.0, "space_before_pt": 12.0, "space_after_pt": 12.0},
    2: {"font_name": "宋体", "font_size_pt": 14.0, "bold": True, "alignment": "left", "line_spacing": 1.25, "left_indent_pt": 0.0, "first_line_indent_pt": 0.0, "space_before_pt": 6.0, "space_after_pt": 6.0},
    3: {"font_name": "宋体", "font_size_pt": 12.0, "bold": True, "alignment": "left", "line_spacing": 1.25, "left_indent_pt": 0.0, "first_line_indent_pt": 0.0, "space_before_pt": 6.0, "space_after_pt": 6.0},
}


def first_or_none(values):
    return values[0] if values else None


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def load_style_metadata(zip_file: ZipFile) -> dict[str, dict[str, int | str | None]]:
    style_map: dict[str, dict[str, int | str | None]] = {}
    if "word/styles.xml" not in zip_file.namelist():
        return style_map

    styles_root = etree.fromstring(zip_file.read("word/styles.xml"))
    for style in styles_root.xpath(".//w:style", namespaces=NS):
        style_id = style.get(f"{{{NS['w']}}}styleId")
        if not style_id:
            continue
        style_name = first_or_none(style.xpath("./w:name/@w:val", namespaces=NS))
        outline_level = first_or_none(style.xpath("./w:pPr/w:outlineLvl/@w:val", namespaces=NS))
        style_map[style_id] = {
            "name": style_name,
            "outline_level": int(outline_level) if outline_level is not None else None,
        }
    return style_map


def extract_paragraph_metadata(docx_path: Path) -> tuple[Document, list[dict[str, object]]]:
    document = Document(docx_path)
    with ZipFile(docx_path) as zip_file:
        style_map = load_style_metadata(zip_file)
        document_root = etree.fromstring(zip_file.read("word/document.xml"))

    body = document_root.find("w:body", namespaces=NS)
    paragraph_elements = [child for child in body if local_name(child.tag) == "p"]

    metadata: list[dict[str, object]] = []
    for index, (paragraph, element) in enumerate(zip(document.paragraphs, paragraph_elements), start=1):
        style_id = first_or_none(element.xpath("./w:pPr/w:pStyle/@w:val", namespaces=NS))
        direct_outline = first_or_none(element.xpath("./w:pPr/w:outlineLvl/@w:val", namespaces=NS))
        style_outline = style_map.get(style_id, {}).get("outline_level")
        alignment = first_or_none(element.xpath("./w:pPr/w:jc/@w:val", namespaces=NS))

        metadata.append(
            {
                "paragraph_index": index,
                "text": normalize_text(paragraph.text),
                "style_id": style_id,
                "style_name": paragraph.style.name if paragraph.style is not None else style_map.get(style_id, {}).get("name", ""),
                "outline_level": int(direct_outline) if direct_outline is not None else style_outline,
                "alignment": alignment,
                "paragraph": paragraph,
            }
        )
    return document, metadata


def is_bold_paragraph(paragraph) -> bool:
    for run in paragraph.runs:
        if run.text.strip() and run.bold is True:
            return True
    return bool(paragraph.style and paragraph.style.font and paragraph.style.font.bold)


def resolve_alignment(meta: dict[str, object]) -> str:
    if meta.get("alignment"):
        return str(meta["alignment"])

    paragraph = meta["paragraph"]
    value = paragraph.alignment
    if value is None and paragraph.style is not None:
        value = paragraph.style.paragraph_format.alignment

    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        WD_ALIGN_PARAGRAPH.DISTRIBUTE: "distribute",
        0: "left",
        1: "center",
        2: "right",
        3: "justify",
        4: "distribute",
    }
    return mapping.get(value, "inherit")


def is_caption_or_catalog(text: str, style_name: str) -> bool:
    if style_name in EXCLUDED_STYLES:
        return True
    return any(pattern.match(text) for pattern in CAPTION_PATTERNS)


def extract_heading_number(text: str) -> tuple[str | None, str | None, int | None]:
    for pattern, level in NUMERIC_HEADING_PATTERNS:
        match = pattern.match(text)
        if match:
            return match.group(1), normalize_text(match.group(2)), level
    return None, None, None


def detect_heading(meta: dict[str, object]) -> dict[str, object]:
    text = str(meta["text"])
    style_name = str(meta["style_name"] or "")
    paragraph = meta["paragraph"]
    outline_level = meta.get("outline_level")

    if not text:
        return {"is_heading": False, "heading_level": None, "heading_number": None, "title_text": "", "detect_basis": "empty"}

    if is_caption_or_catalog(text, style_name):
        return {"is_heading": False, "heading_level": None, "heading_number": None, "title_text": "", "detect_basis": "caption_or_catalog"}

    if text in FRONT_HEADINGS:
        return {
            "is_heading": True,
            "heading_level": 1,
            "heading_number": None,
            "title_text": text,
            "detect_basis": "front_matter",
        }

    heading_number, title_text, number_level = extract_heading_number(text)
    style_hint = style_name in STYLE_HINTS
    short_text = len(text) <= 80 and not text.endswith("。")
    bold_hint = is_bold_paragraph(paragraph)
    centered = resolve_alignment(meta) in {"center", "distribute"}

    if heading_number and number_level is not None and (style_hint or outline_level is not None or short_text):
        return {
            "is_heading": True,
            "heading_level": min(number_level, 3),
            "heading_number": heading_number,
            "title_text": title_text or text,
            "detect_basis": "number_regex",
        }

    if outline_level is not None and style_hint and short_text:
        return {
            "is_heading": True,
            "heading_level": min(int(outline_level) + 1, 3),
            "heading_number": None,
            "title_text": text,
            "detect_basis": "outline_level",
        }

    if style_name == "Title" and short_text:
        return {
            "is_heading": True,
            "heading_level": 3,
            "heading_number": None,
            "title_text": text,
            "detect_basis": "style_title",
        }

    if style_name in {"Subtitle", "1正文"} and short_text:
        return {
            "is_heading": True,
            "heading_level": 1,
            "heading_number": None,
            "title_text": text,
            "detect_basis": "style_subtitle",
        }

    if centered and bold_hint and short_text:
        return {
            "is_heading": True,
            "heading_level": 1,
            "heading_number": None,
            "title_text": text,
            "detect_basis": "layout_fallback",
        }

    return {"is_heading": False, "heading_level": None, "heading_number": None, "title_text": "", "detect_basis": "body"}


def annotate_heading_paths(metadata: list[dict[str, object]]) -> list[dict[str, object]]:
    stack: dict[int, str] = {}
    annotated: list[dict[str, object]] = []

    for item in metadata:
        detection = detect_heading(item)
        row = dict(item)
        row.update(detection)

        if detection["is_heading"]:
            level = int(detection["heading_level"])
            stack[level] = str(item["text"])
            for deeper in range(level + 1, 10):
                stack.pop(deeper, None)
            row["heading_path"] = " > ".join(stack[key] for key in sorted(stack))
        else:
            row["heading_path"] = " > ".join(stack[key] for key in sorted(stack))

        annotated.append(row)

    return annotated


def get_run_font_name(run) -> str | None:
    if run.font.name:
        return run.font.name
    rpr = run._element.rPr
    if rpr is None or rpr.rFonts is None:
        return None
    for attr in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        value = rpr.rFonts.get(qn(attr))
        if value:
            return value
    return None


def get_style_font_name(style) -> str | None:
    current = style
    while current is not None:
        if current.font and current.font.name:
            return current.font.name
        rpr = current.element.rPr
        if rpr is not None and rpr.rFonts is not None:
            for attr in ("w:eastAsia", "w:ascii", "w:hAnsi"):
                value = rpr.rFonts.get(qn(attr))
                if value:
                    return value
        current = current.base_style
    return None


def get_style_font_size_pt(style) -> float | None:
    current = style
    while current is not None:
        if current.font and current.font.size is not None:
            return current.font.size.pt
        current = current.base_style
    return None


def resolve_font_name(paragraph) -> str:
    for run in paragraph.runs:
        if run.text.strip():
            name = get_run_font_name(run)
            if name:
                return name
    if paragraph.style is not None:
        name = get_style_font_name(paragraph.style)
        if name:
            return name
    return "未解析"


def resolve_font_size_pt(paragraph) -> float | None:
    for run in paragraph.runs:
        if run.text.strip() and run.font.size is not None:
            return run.font.size.pt
    if paragraph.style is not None:
        return get_style_font_size_pt(paragraph.style)
    return None


def resolve_style_paragraph_value(paragraph, attr_name: str):
    style = paragraph.style
    while style is not None:
        value = getattr(style.paragraph_format, attr_name)
        if value is not None:
            return value
        style = style.base_style
    return None


def resolve_paragraph_value(paragraph, attr_name: str):
    value = getattr(paragraph.paragraph_format, attr_name)
    if value is not None:
        return value
    return resolve_style_paragraph_value(paragraph, attr_name)


def length_to_pt(value) -> float | None:
    return round(value.pt, 2) if value is not None else None


def describe_line_spacing(value) -> str:
    if value is None:
        return "继承/未设置"
    if hasattr(value, "pt"):
        return f"固定值 {value.pt:.2f} 磅"
    if isinstance(value, (int, float)):
        return f"{float(value):.2f} 倍"
    return str(value)


def apply_run_font(run, font_name: str, font_size_pt: float, bold: bool) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.bold = bold

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rfonts.set(qn(attr), font_name)


def alignment_to_value(name: str):
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
    }
    return mapping[name]


def parse_numeric_sequence(number: str | None) -> list[int] | None:
    if number and re.fullmatch(r"\d+(?:\.\d+){0,2}", number):
        return [int(item) for item in number.split(".")]
    return None


def build_heading_diagnostics(annotated: list[dict[str, object]]) -> list[dict[str, object]]:
    diagnostics: list[dict[str, object]] = []
    expected_counters = [0, 0, 0]

    for row in annotated:
        if not row["is_heading"]:
            continue

        paragraph = row["paragraph"]
        level = int(row["heading_level"])
        scope = "checked" if row["heading_number"] and level in {1, 2, 3} else "informational"
        expected_number = None
        issues: list[str] = []

        if scope == "checked":
            expected_counters[level - 1] += 1
            for deeper in range(level, len(expected_counters)):
                expected_counters[deeper] = 0
            expected_number = ".".join(str(expected_counters[idx]) for idx in range(level))

            if row["heading_number"] != expected_number:
                issues.append(f"编号不连续，建议为 {expected_number}，实际为 {row['heading_number']}")

        font_name = resolve_font_name(paragraph)
        font_size_pt = resolve_font_size_pt(paragraph)
        alignment = resolve_alignment(row)
        left_indent_pt = length_to_pt(resolve_paragraph_value(paragraph, "left_indent"))
        first_line_indent_pt = length_to_pt(resolve_paragraph_value(paragraph, "first_line_indent"))
        line_spacing_value = resolve_paragraph_value(paragraph, "line_spacing")
        space_before_pt = length_to_pt(resolve_paragraph_value(paragraph, "space_before"))
        space_after_pt = length_to_pt(resolve_paragraph_value(paragraph, "space_after"))
        bold_flag = is_bold_paragraph(paragraph)

        rule = LEVEL_RULES.get(level)
        if scope == "checked" and rule:
            if rule["font_name"] not in font_name:
                issues.append(f"字体建议为 {rule['font_name']}，实际为 {font_name}")
            if font_size_pt is None or abs(font_size_pt - rule["font_size_pt"]) > 0.6:
                issues.append(f"字号建议为 {rule['font_size_pt']:.1f} pt，实际为 {font_size_pt if font_size_pt is not None else '未解析'}")
            if alignment != rule["alignment"]:
                issues.append(f"对齐方式建议为 {rule['alignment']}，实际为 {alignment}")
            if left_indent_pt is not None and abs(left_indent_pt - rule["left_indent_pt"]) > 0.6:
                issues.append(f"左缩进建议为 {rule['left_indent_pt']:.1f} pt，实际为 {left_indent_pt:.2f} pt")
            if first_line_indent_pt is not None and abs(first_line_indent_pt - rule["first_line_indent_pt"]) > 0.6:
                issues.append(f"首行缩进建议为 {rule['first_line_indent_pt']:.1f} pt，实际为 {first_line_indent_pt:.2f} pt")
            if isinstance(line_spacing_value, (int, float)) and abs(float(line_spacing_value) - rule["line_spacing"]) > 0.05:
                issues.append(f"行距建议为 {rule['line_spacing']:.2f} 倍，实际为 {float(line_spacing_value):.2f} 倍")
            if rule["bold"] != bold_flag:
                issues.append(f"粗体建议为 {rule['bold']}，实际为 {bold_flag}")

        diagnostics.append(
            {
                "paragraph_index": row["paragraph_index"],
                "display_text": row["text"],
                "title": row["title_text"] or row["text"],
                "detected_level": level,
                "detected_number": row["heading_number"],
                "expected_number": expected_number,
                "style_name": row["style_name"],
                "heading_path": row["heading_path"],
                "detect_basis": row["detect_basis"],
                "font_name": font_name,
                "font_size_pt": font_size_pt,
                "bold": bold_flag,
                "alignment": alignment,
                "left_indent_pt": left_indent_pt,
                "first_line_indent_pt": first_line_indent_pt,
                "line_spacing_desc": describe_line_spacing(line_spacing_value),
                "space_before_pt": space_before_pt,
                "space_after_pt": space_after_pt,
                "scope": scope,
                "issues": issues,
            }
        )

    return diagnostics


def build_normalized_display_text(diagnostic: dict[str, object]) -> str:
    number = diagnostic["detected_number"]
    title = diagnostic["title"]
    if number:
        return f"{number} {title}"
    return str(title)


def normalize_heading_document(diagnostics: list[dict[str, object]], output_path: Path) -> None:
    document = Document(INPUT_DOCX)

    for diagnostic in diagnostics:
        if diagnostic["scope"] != "checked":
            continue

        paragraph = document.paragraphs[diagnostic["paragraph_index"] - 1]
        paragraph.text = build_normalized_display_text(diagnostic)
        if not paragraph.runs:
            paragraph.add_run(paragraph.text)

        run = paragraph.runs[0]
        level = int(diagnostic["detected_level"])
        rule = LEVEL_RULES[level]

        apply_run_font(run, rule["font_name"], rule["font_size_pt"], rule["bold"])
        paragraph.alignment = alignment_to_value(rule["alignment"])
        paragraph.paragraph_format.line_spacing = rule["line_spacing"]
        paragraph.paragraph_format.left_indent = Pt(rule["left_indent_pt"])
        paragraph.paragraph_format.first_line_indent = Pt(rule["first_line_indent_pt"])
        paragraph.paragraph_format.space_before = Pt(rule["space_before_pt"])
        paragraph.paragraph_format.space_after = Pt(rule["space_after_pt"])

    document.save(output_path)


def build_report_text(diagnostics: list[dict[str, object]]) -> str:
    checked = [item for item in diagnostics if item["scope"] == "checked"]
    with_issues = [item for item in checked if item["issues"]]

    lines = [
        "标题诊断报告",
        f"输入文件: {INPUT_DOCX}",
        f"检测到标题数量: {len(diagnostics)}",
        f"纳入章/节/小节规范检查数量: {len(checked)}",
        f"存在问题的标题数量: {len(with_issues)}",
        "",
    ]

    for item in diagnostics:
        lines.append(f"段落 {item['paragraph_index']}: {item['display_text']}")
        lines.append(f"  层级: {item['detected_level']} | 编号: {item['detected_number']} | 建议编号: {item['expected_number']}")
        lines.append(f"  样式: {item['style_name']} | 识别依据: {item['detect_basis']} | 路径: {item['heading_path']}")
        lines.append(
            f"  格式: 字体={item['font_name']}，字号={item['font_size_pt']}，粗体={item['bold']}，"
            f"对齐={item['alignment']}，左缩进={item['left_indent_pt']}pt，首行缩进={item['first_line_indent_pt']}pt，"
            f"行距={item['line_spacing_desc']}，段前={item['space_before_pt']}pt，段后={item['space_after_pt']}pt"
        )
        if item["issues"]:
            for issue in item["issues"]:
                lines.append(f"  问题: {issue}")
        else:
            lines.append("  问题: 无")
        lines.append("")

    return "\n".join(lines)


def main() -> None:
    if not INPUT_DOCX.exists():
        raise FileNotFoundError(f"未找到输入文件: {INPUT_DOCX}")

    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    _, metadata = extract_paragraph_metadata(INPUT_DOCX)
    annotated = annotate_heading_paths(metadata)
    diagnostics = build_heading_diagnostics(annotated)

    json_path = TESTS_DIR / "2.4_标题诊断报告.json"
    txt_path = TESTS_DIR / "2.4_标题诊断报告.txt"
    docx_path = TESTS_DIR / "2.4_标题规范化副本.docx"

    json_path.write_text(json.dumps(diagnostics, ensure_ascii=False, indent=2), encoding="utf-8")
    txt_path.write_text(build_report_text(diagnostics), encoding="utf-8")
    normalize_heading_document(diagnostics, docx_path)

    checked = [item for item in diagnostics if item["scope"] == "checked"]
    issues = [item for item in checked if item["issues"]]

    print(f"[2.4] 输入文档: {INPUT_DOCX}")
    print(f"[2.4] 已输出标题诊断报告 JSON: {json_path}")
    print(f"[2.4] 已输出标题诊断报告 TXT: {txt_path}")
    print(f"[2.4] 已输出标题规范化副本: {docx_path}")
    print(f"[2.4] 纳入检查标题数: {len(checked)}，存在问题标题数: {len(issues)}")


if __name__ == "__main__":
    main()
