from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from _06_docx_styles import prepare_common_styles


TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"
SOURCE_DOCX = TESTS_DIR / "测试文档.docx"
OUTPUT_DOCX = TESTS_DIR / "4.3_添加文档模板结果.docx"
REPORT_JSON = TESTS_DIR / "4.3_模板迁移报告.json"
REPORT_TXT = TESTS_DIR / "4.3_模板迁移报告.txt"


def copy_docx(source: Path, target: Path) -> Path:
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(source, target)
    return target


def apply_template_capabilities(document: Document) -> dict[str, object]:
    prepare_common_styles(document)
    document.settings.odd_and_even_pages_header_footer = True
    document.core_properties.subject = "已注入论文模板能力"
    document.core_properties.keywords = "模板迁移, 论文排版, python-docx"

    for section in document.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(2.54)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)

    first_section = document.sections[0]
    first_section.header.paragraphs[0].text = "论文模板页眉：三维重建技术研究"
    first_section.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    first_section.footer.paragraphs[0].text = "论文模板页脚：测试文档副本"
    first_section.footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if document.paragraphs:
        note = document.paragraphs[0].insert_paragraph_before(
            "本文档副本已应用论文模板能力，包括页面设置、页眉页脚与标题样式体系。"
        )
        note.style = document.styles["ThesisBody"]

    return {
        "source_document": str(SOURCE_DOCX),
        "output_document": str(OUTPUT_DOCX),
        "sections_updated": len(document.sections),
        "styles_injected": ["ThesisTitle1", "ThesisTitle2", "ThesisTitle3", "ThesisBody"],
        "settings": {
            "paper_size": "A4",
            "margins_cm": {"top": 2.54, "bottom": 2.54, "left": 3.18, "right": 2.54},
            "header_footer_distance_cm": 1.5,
            "odd_even_headers": True,
        },
    }


def write_report(report: dict[str, object]) -> None:
    REPORT_JSON.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    lines = [
        "4.3 模板迁移报告",
        f"源文档: {report['source_document']}",
        f"输出文档: {report['output_document']}",
        f"更新节数: {report['sections_updated']}",
        f"注入样式: {', '.join(report['styles_injected'])}",
        f"纸张: {report['settings']['paper_size']}",
        (
            "页边距(cm): "
            f"上{report['settings']['margins_cm']['top']} / 下{report['settings']['margins_cm']['bottom']} / "
            f"左{report['settings']['margins_cm']['left']} / 右{report['settings']['margins_cm']['right']}"
        ),
        f"页眉页脚距离(cm): {report['settings']['header_footer_distance_cm']}",
        f"奇偶页不同页眉页脚: {report['settings']['odd_even_headers']}",
    ]
    REPORT_TXT.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    copy_docx(SOURCE_DOCX, OUTPUT_DOCX)
    document = Document(OUTPUT_DOCX)
    report = apply_template_capabilities(document)
    document.save(OUTPUT_DOCX)
    write_report(report)

    print(f"[4.3] 已生成模板注入结果: {OUTPUT_DOCX}")
    print(f"[4.3] 已输出模板迁移报告: {REPORT_JSON}")


if __name__ == "__main__":
    main()

