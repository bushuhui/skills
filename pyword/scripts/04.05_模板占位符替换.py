from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from _06_docx_styles import prepare_common_styles
from _08_docx_textops import replace_text_everywhere


TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"
TEMPLATE_SOURCE = TESTS_DIR / "4.5_模板内容源.docx"
OUTPUT_DOCX = TESTS_DIR / "4.5_模板内容替换结果.docx"
REPORT_JSON = TESTS_DIR / "4.5_模板替换映射.json"


def build_template(source_path: Path) -> None:
    document = Document()
    prepare_common_styles(document)

    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.header.paragraphs[0].text = "{{SCHOOL_NAME}}"
    section.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    section.footer.paragraphs[0].text = "{{DOCUMENT_TYPE}}"
    section.footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("{{TITLE}}", style="ThesisTitle1")
    document.add_paragraph("作者：{{AUTHOR}}", style="ThesisBody")
    document.add_paragraph("学院：{{SCHOOL_NAME}}", style="ThesisBody")
    document.add_paragraph("摘要", style="ThesisTitle2")
    document.add_paragraph("{{ABSTRACT}}", style="ThesisBody")
    document.add_paragraph("研究内容", style="ThesisTitle2")
    document.add_paragraph("{{BODY_TEXT}}", style="ThesisBody")

    table = document.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "内容"
    table.cell(1, 0).text = "项目名称"
    table.cell(1, 1).text = "{{PROJECT_NAME}}"
    table.cell(2, 0).text = "指导教师"
    table.cell(2, 1).text = "{{SUPERVISOR}}"
    table.cell(3, 0).text = "提交日期"
    table.cell(3, 1).text = "{{SUBMIT_DATE}}"

    source_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(source_path)


def replace_template_content(source: Path, target: Path, replacements: dict[str, str]) -> None:
    shutil.copyfile(source, target)
    document = Document(target)
    replace_text_everywhere(document, list(replacements.items()))
    document.core_properties.title = replacements["{{TITLE}}"]
    document.core_properties.author = replacements["{{AUTHOR}}"]
    document.save(target)


def main() -> None:
    replacements = {
        "{{TITLE}}": "三维重建技术论文模板替换示例",
        "{{AUTHOR}}": "张三",
        "{{SCHOOL_NAME}}": "数字空间信息学院",
        "{{ABSTRACT}}": "本文展示如何使用 python-docx 对模板中的正文、表格、页眉页脚占位符进行统一替换。",
        "{{BODY_TEXT}}": "替换完成后，模板中的结构框架被保留，而具体内容会被新的论文信息覆盖。",
        "{{PROJECT_NAME}}": "校园三维重建平台",
        "{{SUPERVISOR}}": "李教授",
        "{{SUBMIT_DATE}}": "2026-03-31",
        "{{DOCUMENT_TYPE}}": "毕业论文模板替换结果",
    }

    build_template(TEMPLATE_SOURCE)
    replace_template_content(TEMPLATE_SOURCE, OUTPUT_DOCX, replacements)
    REPORT_JSON.write_text(json.dumps(replacements, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"[4.5] 已生成模板源文档: {TEMPLATE_SOURCE}")
    print(f"[4.5] 已生成模板内容替换结果: {OUTPUT_DOCX}")
    print(f"[4.5] 已输出替换映射: {REPORT_JSON}")


if __name__ == "__main__":
    main()

