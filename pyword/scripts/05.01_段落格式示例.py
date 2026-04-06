from __future__ import annotations

from docx import Document
from docx.shared import Pt

import _09_format_presets as PRESETS
from _12_docx_format_helpers import (
    apply_indentation,
    apply_line_spacing,
    apply_spacing_before_after,
    ensure_tests_dir,
    set_default_style,
    set_run_font,
)

DEMO_PARAGRAPHS = [
    {
        "title": "左对齐 + 单倍行距 + 首行缩进 2 字符",
        "text": "本段演示左对齐、单倍行距、段前段后按行设置，并通过字符单位实现首行缩进两个字符，更适合中文正文排版。",
        "alignment": "左对齐",
        "line_spacing_mode": "single",
        "space_before": 0.5,
        "space_after": 0.5,
        "spacing_unit": "lines",
        "first_line_chars": 2,
    },
    {
        "title": "右对齐 + 1.5 倍行距 + 段前后磅值",
        "text": "本段演示右对齐和 1.5 倍行距，同时用磅值精确控制段前段后间距，适用于说明文字或特殊版式。",
        "alignment": "右对齐",
        "line_spacing_mode": "multiple",
        "line_spacing_value": 1.5,
        "space_before": 8,
        "space_after": 10,
        "spacing_unit": "pt",
    },
    {
        "title": "居中 + 固定值 23 磅行距",
        "text": "本段演示居中和固定值 23 磅行距，这类设置常用于封面、题注或强制等高的版式场景。",
        "alignment": "居中",
        "line_spacing_mode": "exact",
        "line_spacing_value": 23,
        "space_before": 0,
        "space_after": 0,
        "spacing_unit": "pt",
    },
    {
        "title": "两端对齐 + 1.25 倍行距 + 左右缩进",
        "text": "本段演示两端对齐、1.25 倍行距以及左右缩进。对于论文正文和较长中文段落，两端对齐通常能得到更规整的版面效果。",
        "alignment": "两端对齐",
        "line_spacing_mode": "multiple",
        "line_spacing_value": 1.25,
        "space_before": 0,
        "space_after": 0.5,
        "spacing_unit": "lines",
        "left_indent_pt": 18,
        "right_indent_pt": 18,
    },
    {
        "title": "分散对齐 + 最小值 20 磅行距 + 悬挂缩进",
        "text": "本段演示分散对齐、最小值行距和悬挂缩进。分散对齐在中文场景中较常见，悬挂缩进则适合参考文献和编号条目。",
        "alignment": "分散对齐",
        "line_spacing_mode": "at_least",
        "line_spacing_value": 20,
        "space_before": 6,
        "space_after": 6,
        "spacing_unit": "pt",
        "left_indent_pt": 24,
        "hanging_chars": 2,
    },
]


def add_demo_paragraph(document: Document, item: dict[str, object]) -> None:
    title = document.add_paragraph()
    title_run = title.add_run(item["title"])
    set_run_font(title_run, "黑体", "四号")
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(3)

    paragraph = document.add_paragraph()
    run = paragraph.add_run(item["text"])
    set_run_font(run, "宋体", "小四")
    paragraph.alignment = PRESETS.resolve_alignment(item["alignment"])
    apply_line_spacing(paragraph, item["line_spacing_mode"], item.get("line_spacing_value"))
    apply_spacing_before_after(
        paragraph,
        before=item.get("space_before", 0),
        after=item.get("space_after", 0),
        unit=item.get("spacing_unit", "pt"),
    )
    apply_indentation(
        paragraph,
        left_indent_pt=item.get("left_indent_pt"),
        right_indent_pt=item.get("right_indent_pt"),
        first_line_indent_pt=item.get("first_line_indent_pt"),
        hanging_indent_pt=item.get("hanging_indent_pt"),
        first_line_chars=item.get("first_line_chars"),
        hanging_chars=item.get("hanging_chars"),
    )


def main() -> None:
    output_dir = ensure_tests_dir(__file__)
    output_path = output_dir / "5.1_段落格式示例.docx"

    document = Document()
    set_default_style(document)

    heading = document.add_paragraph()
    heading_run = heading.add_run("5.1 段落格式示例")
    set_run_font(heading_run, "黑体", "小二")
    heading.alignment = PRESETS.resolve_alignment("居中")

    for item in DEMO_PARAGRAPHS:
        add_demo_paragraph(document, item)

    document.save(output_path)
    print(f"[5.1] 已生成段落格式示例文档: {output_path}")


if __name__ == "__main__":
    main()
