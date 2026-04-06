from __future__ import annotations

from docx import Document
from docx.shared import Pt

import _09_format_presets as PRESETS
from _12_docx_format_helpers import (
    apply_direct_format,
    ensure_tests_dir,
    format_text_fragment,
    set_default_style,
    set_run_border,
    set_run_font,
)


def main() -> None:
    output_dir = ensure_tests_dir(__file__)
    output_path = output_dir / "5.4_文本片段格式示例.docx"

    document = Document()
    set_default_style(document)

    heading = document.add_paragraph()
    heading_run = heading.add_run("5.4 指定文本片段格式示例")
    set_run_font(heading_run, "黑体", "小二")
    heading.alignment = PRESETS.resolve_alignment("居中")
    heading.paragraph_format.space_after = Pt(10)

    intro = document.add_paragraph()
    intro_run = intro.add_run("下面的正文先写入原始 runs，再只对命中的文本片段进行格式设置，不直接改写 paragraph.text。")
    set_run_font(intro_run, "宋体", "小四")
    intro.paragraph_format.space_after = Pt(8)

    paragraph = document.add_paragraph()
    run = paragraph.add_run("三维重建技术")
    apply_direct_format(run, font_name="宋体", font_size="小四", bold=True, color_hex="1F497D")
    run = paragraph.add_run("能够提升数字孪生场景的建模效率，并增强多模态内容的")
    apply_direct_format(run, font_name="宋体", font_size="小四")
    run = paragraph.add_run("表达质量")
    apply_direct_format(run, font_name="宋体", font_size="小四", italic=True)
    run = paragraph.add_run("。")
    apply_direct_format(run, font_name="宋体", font_size="小四")

    format_text_fragment(
        paragraph,
        "数字孪生",
        occurrence=1,
        formatter=lambda target_run: apply_direct_format(
            target_run,
            font_name="仿宋",
            font_size="小四",
            bold=True,
            underline=True,
            color_hex="C00000",
        ),
    )
    format_text_fragment(
        paragraph,
        "表达质量",
        occurrence=1,
        formatter=lambda target_run: (
            apply_direct_format(
                target_run,
                font_name="宋体",
                font_size="小四",
                bold=True,
                color_hex="7030A0",
            ),
            set_run_border(target_run, value="single", color="7030A0", size=12, space=2),
        ),
    )

    document.save(output_path)
    print(f"[5.4] 已生成文本片段格式示例文档: {output_path}")


if __name__ == "__main__":
    main()
