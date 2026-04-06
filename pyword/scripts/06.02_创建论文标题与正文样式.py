from __future__ import annotations

from pathlib import Path

from docx import Document

from _06_docx_styles import SECTION6, apply_style_rule, get_or_add_paragraph_style, set_default_style

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def add_demo(document: Document) -> None:
    document.add_paragraph("论文样式系统示例", style=SECTION6.STYLE_NAMES["title1"])
    document.add_paragraph("研究背景", style=SECTION6.STYLE_NAMES["title2"])
    document.add_paragraph("行业需求", style=SECTION6.STYLE_NAMES["title3"])
    document.add_paragraph(
        "本段正文使用 ThesisBody 样式，演示通过样式统一管理字体、字号、行距、两端对齐和首行缩进两个字符。",
        style=SECTION6.STYLE_NAMES["body"],
    )


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = TESTS_DIR / "6.2_标题正文样式示例.docx"

    document = Document()
    set_default_style(document)
    for style_name, rule in SECTION6.THESIS_STYLE_RULES.items():
        style = get_or_add_paragraph_style(document, style_name)
        apply_style_rule(style, document, rule)

    add_demo(document)
    document.save(output_path)
    print(f"[6.2] 已生成标题正文样式示例文档: {output_path}")


if __name__ == "__main__":
    main()
