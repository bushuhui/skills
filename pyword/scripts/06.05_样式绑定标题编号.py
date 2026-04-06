from __future__ import annotations

from pathlib import Path

from docx import Document

from _05_docx_numbering import apply_numbering, ensure_numbering
from _06_docx_styles import prepare_common_styles, set_default_style
from _10_style_presets import NUMBERING_LEVELS, STYLE_NAMES

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def build_numbering(document: Document) -> int:
    levels = []
    for level_rule in NUMBERING_LEVELS:
        style_id = document.styles[level_rule["style_name"]].style_id
        levels.append(
            {
                "ilvl": level_rule["ilvl"],
                "num_fmt": level_rule["num_fmt"],
                "lvl_text": level_rule["lvl_text"],
                "suff": level_rule["suff"],
                "p_style_id": style_id,
                "left": 0,
                "hanging": 0,
            }
        )
    return ensure_numbering(document, levels)


def add_heading(document: Document, text: str, level: int, num_id: int) -> None:
    paragraph = document.add_paragraph(text, style=STYLE_NAMES[f"title{level}"])
    apply_numbering(paragraph, num_id, level - 1)


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = TESTS_DIR / "6.5_自动标题编号示例.docx"

    document = Document()
    set_default_style(document)
    prepare_common_styles(document)
    num_id = build_numbering(document)

    add_heading(document, "绪论", 1, num_id)
    document.add_paragraph("本章说明研究背景、技术路线和文档自动化排版的总体目标。", style=STYLE_NAMES["body"])
    add_heading(document, "研究背景", 2, num_id)
    document.add_paragraph("随着数字孪生与三维重建应用的扩展，研究文档的层级结构和编号一致性变得越来越重要。", style=STYLE_NAMES["body"])
    add_heading(document, "行业需求", 3, num_id)
    document.add_paragraph("通过样式与编号系统配合，可以避免手动维护 1、1.1、1.1.1 这类编号文本。", style=STYLE_NAMES["body"])
    add_heading(document, "研究目标", 2, num_id)
    document.add_paragraph("目标是建立一套可复用的 Word 样式与自动编号方案，支撑后续目录和格式检查能力。", style=STYLE_NAMES["body"])

    document.save(output_path)
    print(f"[6.5] 已生成自动标题编号示例文档: {output_path}")


if __name__ == "__main__":
    main()
