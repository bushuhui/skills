from __future__ import annotations

from pathlib import Path

from docx import Document

from _06_docx_styles import prepare_common_styles
from _13_docx_title_helpers import ensure_lower_title_styles

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"
OUTPUT_DOCX = TESTS_DIR / "7.1_多级标题插入示例.docx"


def add_heading(document: Document, text: str, level: int) -> None:
    style_map = {
        1: "ThesisTitle1",
        2: "ThesisTitle2",
        3: "ThesisTitle3",
        4: "ThesisTitle4",
        5: "ThesisTitle5",
        6: "ThesisTitle6",
    }
    document.add_paragraph(text, style=style_map[level])


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    prepare_common_styles(document)
    ensure_lower_title_styles(document, ("ThesisTitle4", "ThesisTitle5", "ThesisTitle6"))

    add_heading(document, "三维重建技术研究", 1)
    document.add_paragraph("本示例展示如何在新文档中按一级到六级层次插入标题。", style="ThesisBody")
    add_heading(document, "研究背景", 2)
    document.add_paragraph("二级标题使用四号宋体加粗、居左排列。", style="ThesisBody")
    add_heading(document, "多视角几何建模", 3)
    document.add_paragraph("三级标题之后通常承接方法描述或实验设计。", style="ThesisBody")
    add_heading(document, "(1) 数据采集", 4)
    document.add_paragraph("三级以下标题采用小四号宋体并保持两字符缩进。", style="ThesisBody")
    add_heading(document, "(1) 设备标定", 5)
    document.add_paragraph("更低层标题可继续沿用同一套格式体系。", style="ThesisBody")
    add_heading(document, "(1) 标定板准备", 6)
    document.add_paragraph("该级别一般用于复杂流程拆解，不建议纳入目录。", style="ThesisBody")

    document.save(OUTPUT_DOCX)
    print(f"[7.1] 已生成多级标题插入示例: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
