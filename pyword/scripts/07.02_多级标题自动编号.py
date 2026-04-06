from __future__ import annotations

from pathlib import Path

from docx import Document

from _05_docx_numbering import apply_numbering, ensure_numbering
from _06_docx_styles import prepare_common_styles
from _13_docx_title_helpers import ensure_lower_title_styles

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"
OUTPUT_DOCX = TESTS_DIR / "7.2_自动标题编号示例.docx"


def build_numbering(document: Document) -> int:
    levels = [
        {"ilvl": 0, "num_fmt": "decimal", "lvl_text": "%1", "p_style_id": document.styles["ThesisTitle1"].style_id, "left": 0, "hanging": 0},
        {"ilvl": 1, "num_fmt": "decimal", "lvl_text": "%1.%2", "p_style_id": document.styles["ThesisTitle2"].style_id, "left": 0, "hanging": 0},
        {"ilvl": 2, "num_fmt": "decimal", "lvl_text": "%1.%2.%3", "p_style_id": document.styles["ThesisTitle3"].style_id, "left": 0, "hanging": 0},
        {"ilvl": 3, "num_fmt": "decimal", "lvl_text": "(%4)", "p_style_id": document.styles["ThesisTitle4"].style_id, "left": 720, "hanging": 360},
        {"ilvl": 4, "num_fmt": "decimal", "lvl_text": "(%5)", "p_style_id": document.styles["ThesisTitle5"].style_id, "left": 960, "hanging": 360},
        {"ilvl": 5, "num_fmt": "decimal", "lvl_text": "(%6)", "p_style_id": document.styles["ThesisTitle6"].style_id, "left": 1200, "hanging": 360},
    ]
    return ensure_numbering(document, levels)


def add_numbered_heading(document: Document, num_id: int, text: str, level: int) -> None:
    style_name = {
        1: "ThesisTitle1",
        2: "ThesisTitle2",
        3: "ThesisTitle3",
        4: "ThesisTitle4",
        5: "ThesisTitle5",
        6: "ThesisTitle6",
    }[level]
    paragraph = document.add_paragraph(text, style=style_name)
    apply_numbering(paragraph, num_id, level - 1)


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    prepare_common_styles(document)
    ensure_lower_title_styles(document, ("ThesisTitle4", "ThesisTitle5", "ThesisTitle6"))
    num_id = build_numbering(document)

    add_numbered_heading(document, num_id, "绪论", 1)
    document.add_paragraph("一级标题采用自动编号 1、2、3。", style="ThesisBody")
    add_numbered_heading(document, num_id, "研究背景", 2)
    document.add_paragraph("二级标题自动生成 1.1、1.2 等编号。", style="ThesisBody")
    add_numbered_heading(document, num_id, "数据获取流程", 3)
    document.add_paragraph("三级标题自动生成 1.1.1 形式编号。", style="ThesisBody")
    add_numbered_heading(document, num_id, "原始影像采集", 4)
    document.add_paragraph("三级以下标题采用 (1)、(2) 的编号形式。", style="ThesisBody")
    add_numbered_heading(document, num_id, "相机参数校准", 5)
    document.add_paragraph("更低层标题继续延续括号编号。", style="ThesisBody")
    add_numbered_heading(document, num_id, "误差评估", 6)
    document.add_paragraph("该脚本演示了一级到六级标题的自动编号能力。", style="ThesisBody")

    document.save(OUTPUT_DOCX)
    print(f"[7.2] 已生成自动标题编号示例: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
