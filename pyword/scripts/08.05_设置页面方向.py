from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.shared import Cm

from _06_docx_styles import prepare_common_styles


TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"
SOURCE_DOCX = TESTS_DIR / "测试文档.docx"
OUTPUT_DOCX = TESTS_DIR / "8.5_页面方向示例.docx"


def copy_docx(source: Path, target: Path) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(source, target)


def main() -> None:
    copy_docx(SOURCE_DOCX, OUTPUT_DOCX)
    document = Document(OUTPUT_DOCX)
    prepare_common_styles(document)

    for section in document.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    landscape_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    landscape_section.orientation = WD_ORIENT.LANDSCAPE
    landscape_section.page_width = Cm(29.7)
    landscape_section.page_height = Cm(21.0)

    document.add_paragraph("横向页面示例", style="ThesisTitle2")
    document.add_paragraph("本节被设置为横向页面，适合放置宽表格或流程图。", style="ThesisBody")
    table = document.add_table(rows=3, cols=6)
    headers = ["阶段", "数据源", "特征数量", "点云规模", "耗时", "备注"]
    for col, value in enumerate(headers):
        table.cell(0, col).text = value
    table.cell(1, 0).text = "采集"
    table.cell(1, 1).text = "多视角影像"
    table.cell(1, 2).text = "256"
    table.cell(1, 3).text = "1200 万点"
    table.cell(1, 4).text = "10 min"
    table.cell(1, 5).text = "适合横向展示"
    table.cell(2, 0).text = "重建"
    table.cell(2, 1).text = "深度网络"
    table.cell(2, 2).text = "512"
    table.cell(2, 3).text = "2300 万点"
    table.cell(2, 4).text = "18 min"
    table.cell(2, 5).text = "结果更完整"

    document.save(OUTPUT_DOCX)
    print(f"[8.5] 已生成页面方向设置示例: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()

