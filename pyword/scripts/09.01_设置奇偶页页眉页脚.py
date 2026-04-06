from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"
SOURCE_DOCX = TESTS_DIR / "测试文档.docx"
OUTPUT_DOCX = TESTS_DIR / "9.1_奇偶页页眉页脚示例.docx"


def copy_docx(source: Path, target: Path) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(source, target)


def clear_paragraph(paragraph) -> None:
    element = paragraph._p
    for child in list(element):
        if child.tag != qn("w:pPr"):
            element.remove(child)


def set_text(paragraph, text: str, alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = alignment
    paragraph.add_run(text)


def ensure_paragraph(container):
    if container.paragraphs:
        return container.paragraphs[0]
    return container.add_paragraph()


def main() -> None:
    copy_docx(SOURCE_DOCX, OUTPUT_DOCX)
    document = Document(OUTPUT_DOCX)
    document.settings.odd_and_even_pages_header_footer = True

    for index, section in enumerate(document.sections, start=1):
        section.header.is_linked_to_previous = False
        section.even_page_header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        section.even_page_footer.is_linked_to_previous = False

        set_text(ensure_paragraph(section.header), f"奇数页页眉 | 第 {index} 节 | 三维重建技术研究")
        set_text(ensure_paragraph(section.even_page_header), f"偶数页页眉 | 第 {index} 节 | 三维重建技术研究")
        set_text(ensure_paragraph(section.footer), f"奇数页页脚 | 第 {index} 节")
        set_text(ensure_paragraph(section.even_page_footer), f"偶数页页脚 | 第 {index} 节")

    document.save(OUTPUT_DOCX)
    print(f"[9.1] 已生成奇偶页不同页眉页脚示例: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
