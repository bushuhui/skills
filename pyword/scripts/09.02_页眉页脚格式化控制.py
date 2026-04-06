from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from _03_docx_fields import enable_update_fields_on_open


TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"
SOURCE_DOCX = TESTS_DIR / "测试文档.docx"
OUTPUT_DOCX = TESTS_DIR / "9.2_页眉页脚格式控制示例.docx"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


def copy_docx(source: Path, target: Path) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(source, target)


def clear_paragraph(paragraph) -> None:
    element = paragraph._p
    for child in list(element):
        if child.tag != qn("w:pPr"):
            element.remove(child)


def set_run_font(run, east_asia: str = "宋体", ascii_name: str = "Times New Roman", size_pt: float = 10.5, bold: bool = False) -> None:
    run.font.name = ascii_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), ascii_name)
    r_fonts.set(qn("w:hAnsi"), ascii_name)
    r_fonts.set(qn("w:cs"), ascii_name)


def add_field(paragraph, instruction: str, placeholder: str) -> None:
    begin = paragraph.add_run()
    set_run_font(begin)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    fld_char.set(qn("w:dirty"), "true")
    begin._r.append(fld_char)

    instr_run = paragraph.add_run()
    set_run_font(instr_run)
    instr = OxmlElement("w:instrText")
    instr.set(XML_SPACE, "preserve")
    instr.text = instruction
    instr_run._r.append(instr)

    separate = paragraph.add_run()
    set_run_font(separate)
    separate_char = OxmlElement("w:fldChar")
    separate_char.set(qn("w:fldCharType"), "separate")
    separate._r.append(separate_char)

    placeholder_run = paragraph.add_run(placeholder)
    set_run_font(placeholder_run)

    end = paragraph.add_run()
    set_run_font(end)
    end_char = OxmlElement("w:fldChar")
    end_char.set(qn("w:fldCharType"), "end")
    end._r.append(end_char)


def configure_text_paragraph(paragraph, text: str, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold: bool = False) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size_pt=10.5, bold=bold)


def configure_footer_with_page_fields(paragraph, label: str) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    prefix = paragraph.add_run(f"{label} 第 ")
    set_run_font(prefix, size_pt=10.5)
    add_field(paragraph, " PAGE ", "1")
    middle = paragraph.add_run(" 页 / 共 ")
    set_run_font(middle, size_pt=10.5)
    add_field(paragraph, " NUMPAGES ", "10")
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size_pt=10.5)


def main() -> None:
    copy_docx(SOURCE_DOCX, OUTPUT_DOCX)
    document = Document(OUTPUT_DOCX)
    enable_update_fields_on_open(document)
    document.settings.odd_and_even_pages_header_footer = True

    section = document.sections[0]
    section.different_first_page_header_footer = True
    section.header.is_linked_to_previous = False
    section.even_page_header.is_linked_to_previous = False
    section.first_page_header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.even_page_footer.is_linked_to_previous = False
    section.first_page_footer.is_linked_to_previous = False
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)

    configure_text_paragraph(section.first_page_header.paragraphs[0], "毕业论文", bold=True)
    configure_text_paragraph(section.header.paragraphs[0], "三维重建技术研究", bold=True)
    configure_text_paragraph(section.even_page_header.paragraphs[0], "THESIS ON 3D RECONSTRUCTION", bold=True)

    configure_footer_with_page_fields(section.first_page_footer.paragraphs[0], "首页")
    configure_footer_with_page_fields(section.footer.paragraphs[0], "奇数页")
    configure_footer_with_page_fields(section.even_page_footer.paragraphs[0], "偶数页")

    document.save(OUTPUT_DOCX)
    print(f"[9.2] 已生成页眉页脚格式控制示例: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()

