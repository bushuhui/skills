from __future__ import annotations

from pathlib import Path

from docx import Document

from _03_docx_fields import add_toc_field, enable_update_fields_on_open
from _06_docx_styles import prepare_common_styles

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = TESTS_DIR / "10.1_中文目录示例.docx"

    document = Document()
    prepare_common_styles(document)
    enable_update_fields_on_open(document)

    document.add_paragraph("目录", style="ToCTitleCn")
    toc_paragraph = document.add_paragraph()
    add_toc_field(toc_paragraph)
    document.add_page_break()

    document.add_paragraph("绪论", style="ThesisTitle1")
    document.add_paragraph("本章说明中文目录的插入方式以及目录字段在 Word 中的更新机制。", style="ThesisBody")
    document.add_paragraph("研究背景", style="ThesisTitle2")
    document.add_paragraph("目录字段依赖标题样式或 outlineLvl 生成条目，因此前面的样式系统可以直接复用。", style="ThesisBody")
    document.add_paragraph("实现路径", style="ThesisTitle3")
    document.add_paragraph("程序通过插入 TOC 域而不是手工拼接目录文本，实现更接近 Word 原生行为的目录。", style="ThesisBody")

    document.save(output_path)
    print(f"[10.1] 已生成中文目录示例文档: {output_path}")


if __name__ == "__main__":
    main()

