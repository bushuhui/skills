from __future__ import annotations

from pathlib import Path

from docx import Document

from _03_docx_fields import add_toc_field, enable_update_fields_on_open
from _06_docx_styles import prepare_common_styles

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = TESTS_DIR / "10.2_英文目录示例.docx"

    document = Document()
    prepare_common_styles(document)
    enable_update_fields_on_open(document)

    document.add_paragraph("Table of Contents", style="ToCTitleEn")
    toc_paragraph = document.add_paragraph()
    add_toc_field(toc_paragraph)
    document.add_page_break()

    document.add_paragraph("Introduction", style="ThesisTitle1")
    document.add_paragraph("This chapter demonstrates how an English table of contents can be inserted using a Word field.", style="ThesisBody")
    document.add_paragraph("Research Background", style="ThesisTitle2")
    document.add_paragraph("The TOC field reuses the heading styles defined in the style system.", style="ThesisBody")
    document.add_paragraph("Implementation Notes", style="ThesisTitle3")
    document.add_paragraph("The field will update automatically when the document is opened in Microsoft Word.", style="ThesisBody")

    document.save(output_path)
    print(f"[10.2] 已生成英文目录示例文档: {output_path}")


if __name__ == "__main__":
    main()

