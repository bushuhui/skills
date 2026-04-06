from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from _03_docx_fields import add_toc_field, enable_update_fields_on_open
from _06_docx_styles import prepare_common_styles, set_rfonts, SECTION6, get_or_add_paragraph_style

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def format_toc_styles(document) -> None:
    toc_specs = {
        "TOC 1": ("宋体", 12.0, 0),
        "TOC 2": ("宋体", 10.5, 12),
        "TOC 3": ("宋体", 10.5, 24),
    }
    for style_name, (font_name, font_size_pt, left_indent_pt) in toc_specs.items():
        style = get_or_add_paragraph_style(document, style_name)
        style.base_style = document.styles["Normal"]
        font_spec = SECTION6.resolve_font_family(font_name)
        style.font.name = font_spec["ascii"]
        style.font.size = Pt(font_size_pt)
        set_rfonts(style.element.get_or_add_rPr(), font_spec)
        style.paragraph_format.left_indent = Pt(left_indent_pt)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = TESTS_DIR / "10.3_目录格式调整示例.docx"

    document = Document()
    prepare_common_styles(document)
    enable_update_fields_on_open(document)
    format_toc_styles(document)

    document.add_paragraph("目录格式调整示例", style="ToCTitleCn")
    toc_paragraph = document.add_paragraph()
    add_toc_field(toc_paragraph)
    document.add_page_break()

    document.add_paragraph("总体设计", style="ThesisTitle1")
    document.add_paragraph("目录格式通过修改 TOC 1、TOC 2、TOC 3 等内置样式来统一控制。", style="ThesisBody")
    document.add_paragraph("实现方式", style="ThesisTitle2")
    document.add_paragraph("这样更新目录后格式仍然保持一致，而不是被重新生成的目录覆盖。", style="ThesisBody")
    document.add_paragraph("适用场景", style="ThesisTitle3")
    document.add_paragraph("适用于论文、技术报告以及需要重复更新目录的文档。", style="ThesisBody")

    document.save(output_path)
    print(f"[10.3] 已生成目录格式调整示例文档: {output_path}")


if __name__ == "__main__":
    main()

