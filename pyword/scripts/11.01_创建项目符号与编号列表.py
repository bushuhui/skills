from __future__ import annotations

from pathlib import Path

from docx import Document

from _05_docx_numbering import apply_numbering, ensure_numbering
from _06_docx_styles import prepare_common_styles

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = TESTS_DIR / "11.1_项目符号和编号列表示例.docx"

    document = Document()
    prepare_common_styles(document)

    bullet_num_id = ensure_numbering(
        document,
        [{"ilvl": 0, "num_fmt": "bullet", "lvl_text": "•", "left": 360, "hanging": 360, "bullet_font": "Symbol"}],
    )
    number_num_id = ensure_numbering(
        document,
        [{"ilvl": 0, "num_fmt": "decimal", "lvl_text": "%1.", "left": 360, "hanging": 360}],
    )

    document.add_paragraph("列表示例", style="ThesisTitle1")
    document.add_paragraph("下面演示项目符号列表和编号列表，均采用 Word 原生 numbering.xml 结构。", style="ThesisBody")

    for text in ("数据采集", "特征提取", "稠密重建"):
        paragraph = document.add_paragraph(text, style="ListBody")
        apply_numbering(paragraph, bullet_num_id, 0)

    for text in ("定义问题", "建立模型", "输出结果"):
        paragraph = document.add_paragraph(text, style="ListBody")
        apply_numbering(paragraph, number_num_id, 0)

    document.save(output_path)
    print(f"[11.1] 已生成项目符号和编号列表示例文档: {output_path}")


if __name__ == "__main__":
    main()

