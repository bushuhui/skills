from __future__ import annotations

from pathlib import Path

from docx import Document

from _02_docx_captions import add_caption_paragraph
from _06_docx_styles import prepare_common_styles
from _07_docx_tables import format_cell_text
from _11_docx_enhance_presets import SAMPLE_TABLE_DATA

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = TESTS_DIR / "12.1_创建和插入表格示例.docx"

    document = Document()
    prepare_common_styles(document)

    document.add_paragraph("表格创建与插入示例", style="ThesisTitle1")
    document.add_paragraph("本示例创建一张带自定义数据的表格，并在插入表格前保留一空行。", style="ThesisBody")
    document.add_paragraph("")
    add_caption_paragraph(document, "表", "表", "基础实验数据表", "TableCaption", "tbl_basic_insert")

    table = document.add_table(rows=len(SAMPLE_TABLE_DATA), cols=len(SAMPLE_TABLE_DATA[0]))
    for row_idx, row_data in enumerate(SAMPLE_TABLE_DATA):
        for col_idx, value in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = value
            format_cell_text(cell, bold=(row_idx == 0))

    document.save(output_path)
    print(f"[12.1] 已生成创建和插入表格示例文档: {output_path}")


if __name__ == "__main__":
    main()

