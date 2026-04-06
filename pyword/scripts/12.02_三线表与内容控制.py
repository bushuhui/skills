from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from _02_docx_captions import add_caption_paragraph
from _06_docx_styles import prepare_common_styles
from _07_docx_tables import apply_three_line_table, format_cell_text
from _11_docx_enhance_presets import SAMPLE_TABLE_DATA

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = TESTS_DIR / "12.2_三线表和表格内容示例.docx"

    document = Document()
    prepare_common_styles(document)

    document.add_paragraph("三线表与表格内容控制示例", style="ThesisTitle1")
    document.add_paragraph("本示例演示三线表样式、单元格文字格式以及表格内容修改。", style="ThesisBody")
    document.add_paragraph("")
    add_caption_paragraph(document, "表", "表", "三线表实验结果", "TableCaption", "tbl_three_line")

    table = document.add_table(rows=len(SAMPLE_TABLE_DATA), cols=len(SAMPLE_TABLE_DATA[0]))
    for row_idx, row_data in enumerate(SAMPLE_TABLE_DATA):
        for col_idx, value in enumerate(row_data):
            if row_idx == len(SAMPLE_TABLE_DATA) - 1 and col_idx == 3:
                value = "修改后：已补充性能说明"
            cell = table.cell(row_idx, col_idx)
            cell.text = value
            format_cell_text(
                cell,
                bold=(row_idx == 0),
                alignment=WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT,
            )

    apply_three_line_table(table)

    document.save(output_path)
    print(f"[12.2] 已生成三线表和表格内容示例文档: {output_path}")


if __name__ == "__main__":
    main()

