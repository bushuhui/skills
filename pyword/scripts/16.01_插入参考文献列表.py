from __future__ import annotations

from pathlib import Path

from docx import Document

from _01_docx_bookmarks import add_bookmark_around_run
from _06_docx_styles import prepare_common_styles
from _11_docx_enhance_presets import SAMPLE_BIBLIOGRAPHY

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = TESTS_DIR / "16.1_参考文献插入示例.docx"

    document = Document()
    prepare_common_styles(document)

    document.add_paragraph("参考文献", style="ThesisTitle1")
    for entry in SAMPLE_BIBLIOGRAPHY:
        paragraph = document.add_paragraph(style="BibliographyEntry")
        number_run = paragraph.add_run(f"[{entry['number']}] ")
        add_bookmark_around_run(document, number_run, entry["bookmark"])
        paragraph.add_run(entry["text"])

    document.save(output_path)
    print(f"[16.1] 已生成参考文献插入示例文档: {output_path}")


if __name__ == "__main__":
    main()

