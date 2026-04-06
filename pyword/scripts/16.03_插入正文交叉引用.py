from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from _01_docx_bookmarks import add_bookmark_to_paragraph
from _02_docx_captions import add_caption_paragraph
from _03_docx_fields import add_pageref_field, add_ref_field, enable_update_fields_on_open
from _06_docx_styles import prepare_common_styles
from _11_docx_enhance_presets import SAMPLE_IMAGE_1

TESTS_DIR = Path(__file__).resolve().parent.parent / "tests"


def main() -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = TESTS_DIR / "16.3_正文交叉引用示例.docx"

    document = Document()
    prepare_common_styles(document)
    enable_update_fields_on_open(document)

    document.add_paragraph("正文交叉引用示例", style="ThesisTitle1")
    intro = document.add_paragraph(style="ThesisBody")
    intro.add_run("详见章节 ")
    add_ref_field(intro, "sec_experiment_setup", "实验设置")
    intro.add_run("（第 ")
    add_pageref_field(intro, "sec_experiment_setup", "1")
    intro.add_run(" 页），处理流程如 ")
    add_ref_field(intro, "fig_pipeline_demo", "图 1")
    intro.add_run(" 所示。")

    heading = document.add_paragraph("实验设置", style="ThesisTitle1")
    add_bookmark_to_paragraph(document, heading, "sec_experiment_setup")
    document.add_paragraph("本节说明实验环境、流程与关键参数设置。", style="ThesisBody")

    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(SAMPLE_IMAGE_1), width=Cm(8))
    add_caption_paragraph(document, "图", "图", "处理流程示意图", "FigureCaption", "fig_pipeline_demo")

    document.save(output_path)
    print(f"[16.3] 已生成正文交叉引用示例文档: {output_path}")


if __name__ == "__main__":
    main()

