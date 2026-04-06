from __future__ import annotations

import json
import shutil
from datetime import datetime, timezone
from pathlib import Path

from docx import Document

from _06_docx_styles import prepare_common_styles
from _14_docx_review_demo import TESTS_DIR
from _16_docx_revisions import (
    build_package_map,
    ensure_track_revisions_setting,
    find_top_level_paragraph_by_text,
    insert_empty_paragraph_after,
    next_revision_id,
    paragraph_diff_segments,
    parse_document_root,
    process_revisions_docx,
    replace_paragraph_segments,
    update_document_root,
    write_package_map,
)

SOURCE_DOCX = TESTS_DIR / "19.5_修订源文档.docx"
TRACKED_DOCX = TESTS_DIR / "19.5_带修订文档.docx"
ACCEPTED_DOCX = TESTS_DIR / "19.5_接受修订结果.docx"
REJECTED_DOCX = TESTS_DIR / "19.5_拒绝修订结果.docx"
REPORT_JSON = TESTS_DIR / "19.5_修订处理报告.json"

AUTHOR = "PyWord"

INSERT_OLD = "多视角影像采集能够为三维重建提供稳定的数据基础。"
INSERT_NEW = "多视角影像采集能够为三维重建提供更高分辨率且稳定的数据基础。"
REPLACE_OLD = "点云配准与网格重建是实验流程中的关键环节。"
REPLACE_NEW = "点云配准与网格重建是实验流程中的核心环节。"
MOVE_TEXT = "传统手工配准流程仍然具有一定参考价值。"
DELETE_TEXT = "过度依赖人工检查会明显增加项目周期。"
ANCHOR_TEXT = "自动化流程能够显著缩短建模周期。"


def build_source_document(path: Path) -> None:
    TESTS_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    prepare_common_styles(document)

    document.add_paragraph("接受与拒绝修订示例", style="ThesisTitle1")
    document.add_paragraph("本示例文档用于演示插入、删除、替换与移动修订的接受和拒绝。", style="ThesisBody")
    document.add_paragraph("研究过程", style="ThesisTitle2")
    document.add_paragraph(INSERT_OLD, style="ThesisBody")
    document.add_paragraph(REPLACE_OLD, style="ThesisBody")
    document.add_paragraph(MOVE_TEXT, style="ThesisBody")
    document.add_paragraph(DELETE_TEXT, style="ThesisBody")
    document.add_paragraph(ANCHOR_TEXT, style="ThesisBody")
    document.add_paragraph("实验结果表明该流程具有良好的精度与稳定性。", style="ThesisBody")
    document.save(path)


def build_tracked_document(source_path: Path, target_path: Path) -> dict[str, object]:
    shutil.copyfile(source_path, target_path)
    package_map = build_package_map(target_path)
    ensure_track_revisions_setting(package_map, enabled=True)
    document_root = parse_document_root(package_map)

    timestamp = datetime.now(timezone.utc).replace(microsecond=0).isoformat()
    next_id_value = next_revision_id(document_root)

    insert_paragraph = find_top_level_paragraph_by_text(document_root, INSERT_OLD)
    next_id_value = replace_paragraph_segments(
        insert_paragraph,
        paragraph_diff_segments(INSERT_OLD, INSERT_NEW),
        AUTHOR,
        next_id_value,
        timestamp,
    )

    replace_paragraph = find_top_level_paragraph_by_text(document_root, REPLACE_OLD)
    next_id_value = replace_paragraph_segments(
        replace_paragraph,
        paragraph_diff_segments(REPLACE_OLD, REPLACE_NEW),
        AUTHOR,
        next_id_value,
        timestamp,
    )

    move_from_paragraph = find_top_level_paragraph_by_text(document_root, MOVE_TEXT)
    next_id_value = replace_paragraph_segments(
        move_from_paragraph,
        [("moveFrom", MOVE_TEXT)],
        AUTHOR,
        next_id_value,
        timestamp,
    )

    delete_paragraph = find_top_level_paragraph_by_text(document_root, DELETE_TEXT)
    next_id_value = replace_paragraph_segments(
        delete_paragraph,
        [("delete", DELETE_TEXT)],
        AUTHOR,
        next_id_value,
        timestamp,
    )

    anchor_paragraph = find_top_level_paragraph_by_text(document_root, ANCHOR_TEXT)
    move_to_paragraph = insert_empty_paragraph_after(anchor_paragraph, clone_properties_from=anchor_paragraph)
    next_id_value = replace_paragraph_segments(
        move_to_paragraph,
        [("moveTo", MOVE_TEXT)],
        AUTHOR,
        next_id_value,
        timestamp,
    )

    update_document_root(package_map, document_root)
    write_package_map(package_map, target_path)

    return {
        "insertions": 2,
        "deletions": 2,
        "move_from": 1,
        "move_to": 1,
        "last_revision_id": next_id_value - 1,
    }


def main() -> None:
    build_source_document(SOURCE_DOCX)
    tracked_counts = build_tracked_document(SOURCE_DOCX, TRACKED_DOCX)
    accepted_counts = process_revisions_docx(TRACKED_DOCX, ACCEPTED_DOCX, accept=True)
    rejected_counts = process_revisions_docx(TRACKED_DOCX, REJECTED_DOCX, accept=False)

    report = {
        "source_document": str(SOURCE_DOCX),
        "tracked_document": str(TRACKED_DOCX),
        "accepted_document": str(ACCEPTED_DOCX),
        "rejected_document": str(REJECTED_DOCX),
        "tracked_change_summary": tracked_counts,
        "accepted_summary": accepted_counts,
        "rejected_summary": rejected_counts,
        "changes": [
            {"type": "insert", "from": INSERT_OLD, "to": INSERT_NEW},
            {"type": "replace", "from": REPLACE_OLD, "to": REPLACE_NEW},
            {"type": "move", "text": MOVE_TEXT, "after": ANCHOR_TEXT},
            {"type": "delete", "text": DELETE_TEXT},
        ],
    }
    REPORT_JSON.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"[19.5] 已生成修订源文档: {SOURCE_DOCX}")
    print(f"[19.5] 已生成带修订文档: {TRACKED_DOCX}")
    print(f"[19.5] 已生成接受修订结果: {ACCEPTED_DOCX}")
    print(f"[19.5] 已生成拒绝修订结果: {REJECTED_DOCX}")
    print(f"[19.5] 已输出修订处理报告: {REPORT_JSON}")


if __name__ == "__main__":
    main()
