from __future__ import annotations

import importlib.util
from pathlib import Path

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from _11_docx_enhance_presets import EXTRA_STYLE_RULES

CURRENT_DIR = Path(__file__).resolve().parent


def _load_section6_presets():
    module_path = CURRENT_DIR / "_10_style_presets.py"
    spec = importlib.util.spec_from_file_location("section6_style_presets", module_path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


SECTION6 = _load_section6_presets()


def get_style(document, style_name: str):
    for style in document.styles:
        if style.name == style_name:
            return style
    raise KeyError(style_name)


def get_or_add_paragraph_style(document, style_name: str):
    try:
        return get_style(document, style_name)
    except KeyError:
        return document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)


def set_rfonts(r_pr, font_spec: dict[str, str]) -> None:
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key, value in font_spec.items():
        r_fonts.set(qn(f"w:{key}"), value)


def set_line_spacing(paragraph_format, mode: str, value: float | None = None) -> None:
    if mode == "multiple":
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph_format.line_spacing = float(value)
    elif mode == "exact":
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph_format.line_spacing = Pt(float(value))
    elif mode == "at_least":
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        paragraph_format.line_spacing = Pt(float(value))
    else:
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph_format.line_spacing = 1.0


def set_line_based_spacing(style, before_lines: float | None, after_lines: float | None) -> None:
    if before_lines is None and after_lines is None:
        return
    p_pr = style._element.get_or_add_pPr()
    spacing = p_pr.spacing
    if spacing is None:
        spacing = p_pr.get_or_add_spacing()
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    if before_lines is not None:
        spacing.set(qn("w:beforeLines"), str(int(round(before_lines * SECTION6.LINE_UNIT_SCALE))))
    if after_lines is not None:
        spacing.set(qn("w:afterLines"), str(int(round(after_lines * SECTION6.LINE_UNIT_SCALE))))


def set_char_indentation(style, first_line_chars: float | None = None, hanging_chars: float | None = None) -> None:
    if first_line_chars is None and hanging_chars is None:
        return
    p_pr = style._element.get_or_add_pPr()
    ind = p_pr.ind
    if ind is None:
        ind = p_pr.get_or_add_ind()
    style.paragraph_format.first_line_indent = Pt(0)
    if first_line_chars is not None:
        ind.set(qn("w:firstLineChars"), str(int(round(first_line_chars * SECTION6.CHAR_UNIT_SCALE))))
    if hanging_chars is not None:
        ind.set(qn("w:hangingChars"), str(int(round(hanging_chars * SECTION6.CHAR_UNIT_SCALE))))
    for attr in ("w:firstLine", "w:hanging"):
        if ind.get(qn(attr)) is not None:
            del ind.attrib[qn(attr)]


def set_outline_level(style, level: int | None) -> None:
    p_pr = style._element.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if level is None:
        if outline is not None:
            p_pr.remove(outline)
        return
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level))


def set_keep_rules(style, keep_next: bool = False, keep_lines: bool = False) -> None:
    p_pr = style._element.get_or_add_pPr()
    for tag, enabled in (("w:keepNext", keep_next), ("w:keepLines", keep_lines)):
        existing = p_pr.find(qn(tag))
        if enabled:
            if existing is None:
                p_pr.append(OxmlElement(tag))
        elif existing is not None:
            p_pr.remove(existing)


def apply_style_rule(style, document, rule: dict[str, object]) -> None:
    style.base_style = document.styles[rule["base_style"]]
    style.hidden = False
    style.quick_style = True
    style.unhide_when_used = True
    style.priority = int(rule.get("ui_priority", 20))

    next_style_name = rule.get("next_style")
    if next_style_name:
        next_style = get_or_add_paragraph_style(document, str(next_style_name))
        style._element.get_or_add_next().val = next_style.style_id

    font_spec = SECTION6.resolve_font_family(str(rule["font_name"]))
    style.font.name = font_spec["ascii"]
    style.font.size = Pt(SECTION6.resolve_font_size_pt(rule["font_size"]))
    style.font.bold = bool(rule.get("bold", False))
    style.font.italic = bool(rule.get("italic", False))
    set_rfonts(style._element.get_or_add_rPr(), font_spec)

    paragraph_format = style.paragraph_format
    paragraph_format.alignment = SECTION6.resolve_alignment(str(rule["alignment"]))
    set_line_spacing(
        paragraph_format,
        str(rule.get("line_spacing_mode", "single")),
        float(rule["line_spacing_value"]) if rule.get("line_spacing_value") is not None else None,
    )
    if rule.get("space_before_pt") is not None:
        paragraph_format.space_before = Pt(float(rule["space_before_pt"]))
    if rule.get("space_after_pt") is not None:
        paragraph_format.space_after = Pt(float(rule["space_after_pt"]))
    if rule.get("left_indent_pt") is not None:
        paragraph_format.left_indent = Pt(float(rule["left_indent_pt"]))
    if rule.get("right_indent_pt") is not None:
        paragraph_format.right_indent = Pt(float(rule["right_indent_pt"]))

    set_line_based_spacing(style, rule.get("space_before_lines"), rule.get("space_after_lines"))
    set_char_indentation(style, rule.get("first_line_chars"), rule.get("hanging_chars"))
    set_outline_level(style, rule.get("outline_level"))
    set_keep_rules(style, bool(rule.get("keep_next", False)), bool(rule.get("keep_lines", False)))


def set_default_style(document) -> None:
    normal = document.styles["Normal"]
    font_spec = SECTION6.resolve_font_family("宋体")
    normal.font.name = font_spec["ascii"]
    normal.font.size = Pt(SECTION6.resolve_font_size_pt("小四"))
    set_rfonts(normal.element.get_or_add_rPr(), font_spec)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)


def prepare_common_styles(document) -> None:
    set_default_style(document)
    for style_name, rule in SECTION6.THESIS_STYLE_RULES.items():
        style = get_or_add_paragraph_style(document, style_name)
        apply_style_rule(style, document, rule)
    for style_name, rule in EXTRA_STYLE_RULES.items():
        style = get_or_add_paragraph_style(document, style_name)
        apply_style_rule(style, document, rule)
